"""
==============================================================================
  Legal Document PII Anonymizer & Restorer  —  GUI Edition  (v3.0)
  -----------------------------------------------------------------
  Full Hebrew + English support.

  Hebrew capabilities:
    * xx_ent_wiki_sm  — multilingual spaCy NER (PERSON / LOC / ORG in Hebrew)
    * HebrewNameRecognizer  — dictionary of 700+ Israeli first/last names
    * HebrewLocationRecognizer — 150+ Israeli cities, regions, legal terms
    * Israeli ID (Teudat Zehut) with Luhn-10 checksum
    * Israeli phone numbers (+972 / 05x / 07x / 0x formats)
    * Israeli IBAN (IL prefix)
    * Hebrew date patterns (DD/MM/YYYY, DD.MM.YYYY)
    * Auto language detection per document
    * RTL-aware text display

  Core engine: Microsoft Presidio + spaCy (en + xx multilingual)
==============================================================================
"""

import json
import logging
import logging.handlers
import os
import re
import sys
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from typing import Dict, List, Optional, Tuple
import urllib.request
import urllib.error
import ssl
import tempfile
import shutil
import subprocess
import datetime

# ---------------------------------------------------------------------------
#  Version — bump this for every release
# ---------------------------------------------------------------------------
APP_VERSION = "0.5.4"
UPDATE_MANIFEST_URL = (
    "https://github.com/17Shaun01/pii-processor-build/"
    "releases/latest/download/version.json"
)

# ---------------------------------------------------------------------------
#  SSL context helper — uses certifi CA bundle when available so that
#  HTTPS requests work correctly inside a PyInstaller bundle on Windows
# ---------------------------------------------------------------------------

def _make_ssl_context() -> ssl.SSLContext:
    """Return an SSL context that trusts the system/certifi CA bundle."""
    try:
        import certifi
        ctx = ssl.create_default_context(cafile=certifi.where())
        return ctx
    except ImportError:
        pass
    try:
        ctx = ssl.create_default_context()
        return ctx
    except Exception:
        pass
    # Last resort: unverified (not ideal but better than crashing)
    ctx = ssl._create_unverified_context()
    return ctx


# ---------------------------------------------------------------------------
#  Logging setup
# ---------------------------------------------------------------------------

def _get_log_path() -> str:
    """Return the log file path next to the executable (or cwd in dev mode)."""
    if getattr(sys, 'frozen', False):
        base = os.path.dirname(sys.executable)
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, "pii_processor.log")


LOG_PATH = _get_log_path()

# Root logger for the app — all modules use getLogger(__name__) or this
logger = logging.getLogger("pii_processor")
logger.setLevel(logging.DEBUG)  # capture everything; handlers filter level

# Rotating file handler — max 2 MB, keep 3 backups
_file_handler = logging.handlers.RotatingFileHandler(
    LOG_PATH, maxBytes=2 * 1024 * 1024, backupCount=3, encoding="utf-8"
)
_file_handler.setLevel(logging.DEBUG)
_file_handler.setFormatter(logging.Formatter(
    "%(asctime)s  %(levelname)-8s  %(threadName)-12s  %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
))
logger.addHandler(_file_handler)

# In-memory handler — feeds the live Debug Log tab (max 2000 records)
class _MemHandler(logging.Handler):
    """Thread-safe ring buffer that the Debug Log tab reads from."""
    def __init__(self, capacity: int = 2000):
        super().__init__()
        self._records: list = []
        self._capacity = capacity
        self._lock = threading.Lock()
        self._callbacks: list = []

    def emit(self, record: logging.LogRecord):
        with self._lock:
            self._records.append(record)
            if len(self._records) > self._capacity:
                self._records.pop(0)
        for cb in list(self._callbacks):
            try:
                cb(record)
            except Exception:
                pass

    def get_records(self) -> list:
        with self._lock:
            return list(self._records)

    def add_callback(self, fn):
        self._callbacks.append(fn)

    def remove_callback(self, fn):
        try:
            self._callbacks.remove(fn)
        except ValueError:
            pass


MEM_LOG_HANDLER = _MemHandler(capacity=2000)
MEM_LOG_HANDLER.setLevel(logging.DEBUG)
MEM_LOG_HANDLER.setFormatter(logging.Formatter(
    "%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
))
logger.addHandler(MEM_LOG_HANDLER)

logger.info("=" * 60)
logger.info("PII Processor v%s started", APP_VERSION)
logger.info("Log file: %s", LOG_PATH)
logger.info("Python: %s", sys.version.split()[0])
logger.info("Platform: %s", sys.platform)
logger.info("=" * 60)


try:
    from presidio_analyzer import (
        AnalyzerEngine, RecognizerRegistry,
        PatternRecognizer, Pattern, EntityRecognizer, RecognizerResult,
    )
    from presidio_analyzer.nlp_engine import NlpEngineProvider
except ImportError:
    messagebox.showerror(
        "Missing Dependency",
        "presidio-analyzer is not installed.\n\nRun:\n  pip install presidio-analyzer presidio-anonymizer spacy",
    )
    sys.exit(1)

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_OK = True
except ImportError:
    PDF_OK = False


# ---------------------------------------------------------------------------
#  Hebrew name / location dictionaries  (imported from hebrew_data.py)
# ---------------------------------------------------------------------------

try:
    from hebrew_data import HEBREW_FIRST_NAMES, HEBREW_LAST_NAMES, HEBREW_LOCATIONS
except ImportError:
    # Fallback: empty sets (pattern recognizers still work)
    HEBREW_FIRST_NAMES = set()
    HEBREW_LAST_NAMES = set()
    HEBREW_LOCATIONS = set()


# ---------------------------------------------------------------------------
#  Language detection
# ---------------------------------------------------------------------------

HEBREW_CHARS = re.compile(r'[\u05d0-\u05ea\ufb1d-\ufb4e]')

def detect_language(text: str) -> str:
    total = len(text.replace(" ", "").replace("\n", ""))
    if total == 0:
        return "en"
    hebrew_count = len(HEBREW_CHARS.findall(text))
    return "he" if (hebrew_count / total) > 0.15 else "en"

def is_rtl(text: str) -> bool:
    return detect_language(text) == "he"


# ---------------------------------------------------------------------------
#  Entity configuration
# ---------------------------------------------------------------------------

ENTITY_LABELS: Dict[str, str] = {
    "PERSON":              "PERSON",
    "EMAIL_ADDRESS":       "EMAIL",
    "PHONE_NUMBER":        "PHONE",
    "LOCATION":            "LOCATION",
    "DATE_TIME":           "DATE",
    "IL_ID_NUMBER":        "IL_ID",
    "IL_PHONE":            "IL_PHONE",
    "HE_PERSON":           "PERSON",
    "HE_LOCATION":         "LOCATION",
    "US_SSN":              "SSN",
    "US_PASSPORT":         "PASSPORT",
    "US_DRIVER_LICENSE":   "DRIVER_LICENSE",
    "US_ITIN":             "ITIN",
    "US_BANK_NUMBER":      "BANK_ACCOUNT",
    "CREDIT_CARD":         "CREDIT_CARD",
    "IBAN_CODE":           "IBAN",
    "IP_ADDRESS":          "IP_ADDRESS",
    "URL":                 "URL",
    "CRYPTO":              "CRYPTO_ADDRESS",
    "MEDICAL_LICENSE":     "MEDICAL_LICENSE",
    "NRP":                 "NATIONAL_ID",
}
ALL_ENTITIES = list(ENTITY_LABELS.keys())
DEFAULT_CONFIDENCE = 0.55

HE_ENTITY_NAMES: Dict[str, str] = {
    "PERSON":           "\u05e9\u05dd \u05d0\u05d3\u05dd",
    "EMAIL":            '\u05d3\u05d5\u05d0"\u05dc',
    "PHONE":            "\u05d8\u05dc\u05e4\u05d5\u05df",
    "IL_PHONE":         "\u05d8\u05dc\u05e4\u05d5\u05df \u05d9\u05e9\u05e8\u05d0\u05dc\u05d9",
    "LOCATION":         "\u05de\u05d9\u05e7\u05d5\u05dd",
    "DATE":             "\u05ea\u05d0\u05e8\u05d9\u05da",
    "IL_ID":            "\u05ea.\u05d6.",
    "SSN":              "SSN",
    "PASSPORT":         "\u05d3\u05e8\u05db\u05d5\u05df",
    "DRIVER_LICENSE":   "\u05e8\u05d9\u05e9\u05d9\u05d5\u05df \u05e0\u05d4\u05d9\u05d2\u05d4",
    "IBAN":             "IBAN",
    "BANK_ACCOUNT":     "\u05d7\u05e9\u05d1\u05d5\u05df \u05d1\u05e0\u05e7",
    "CREDIT_CARD":      "\u05db\u05e8\u05d8\u05d9\u05e1 \u05d0\u05e9\u05e8\u05d0\u05d9",
    "NATIONAL_ID":      "\u05de\u05e1\u05e4\u05e8 \u05d6\u05d4\u05d5\u05ea",
}


# ---------------------------------------------------------------------------
#  Colour palette
# ---------------------------------------------------------------------------

DARK_BG      = "#1e1e2e"
PANEL_BG     = "#2a2a3e"
ACCENT       = "#7c6af7"
ACCENT_HOVER = "#9b8df9"
SUCCESS      = "#50fa7b"
WARNING      = "#f1fa8c"
DANGER       = "#ff5555"
TEXT_MAIN    = "#cdd6f4"
TEXT_DIM     = "#6c7086"
BORDER       = "#45475a"
ENTRY_BG     = "#313244"
TAG_PERSON   = "#ff79c6"
TAG_EMAIL    = "#8be9fd"
TAG_PHONE    = "#50fa7b"
TAG_LOCATION = "#ffb86c"
TAG_DATE     = "#f1fa8c"
TAG_ID       = "#ff5555"
TAG_OTHER    = "#bd93f9"

ENTITY_COLOURS = {
    "PERSON":         TAG_PERSON,
    "EMAIL":          TAG_EMAIL,
    "PHONE":          TAG_PHONE,
    "IL_PHONE":       TAG_PHONE,
    "LOCATION":       TAG_LOCATION,
    "DATE":           TAG_DATE,
    "IL_ID":          TAG_ID,
    "SSN":            TAG_ID,
    "PASSPORT":       TAG_ID,
    "DRIVER_LICENSE": TAG_ID,
    "IBAN":           TAG_ID,
    "BANK_ACCOUNT":   TAG_ID,
    "CREDIT_CARD":    TAG_ID,
    "NATIONAL_ID":    TAG_ID,
}


# ---------------------------------------------------------------------------
#  Israeli-specific PII recognizers
# ---------------------------------------------------------------------------

def _luhn10_israeli_id(id_str: str) -> bool:
    id_str = id_str.zfill(9)
    if len(id_str) != 9 or not id_str.isdigit():
        return False
    total = 0
    for i, ch in enumerate(id_str):
        d = int(ch) * ((i % 2) + 1)
        total += d - 9 if d > 9 else d
    return total % 10 == 0


class IsraeliIdRecognizer(PatternRecognizer):
    PATTERNS = [
        Pattern("Israeli ID (plain)",   r"\b\d{9}\b",             0.6),
        Pattern("Israeli ID (dashes)",  r"\b\d{3}-\d{3}-\d{3}\b", 0.7),
    ]
    CONTEXT = [
        "id", "identity", "id number", "identification",
        "teudat", "zehut", "t.z", "tz",
        "\u05ea.\u05d6", "\u05ea\u05e2\u05d5\u05d3\u05ea \u05d6\u05d4\u05d5\u05ea",
        "\u05de\u05e1\u05e4\u05e8 \u05d6\u05d4\u05d5\u05ea", "\u05d6\u05d4\u05d5\u05ea",
    ]

    def __init__(self, lang="he"):
        super().__init__(
            supported_entity="IL_ID_NUMBER",
            patterns=self.PATTERNS,
            context=self.CONTEXT,
            supported_language=lang,
        )

    def validate_result(self, pattern_text: str):
        digits = re.sub(r"[^0-9]", "", pattern_text)
        return _luhn10_israeli_id(digits)


class IsraeliPhoneRecognizer(PatternRecognizer):
    PATTERNS = [
        Pattern("IL Phone (+972)",   r"\+972[-\s]?\d{1,2}[-\s]?\d{3}[-\s]?\d{4}", 0.85),
        Pattern("IL Mobile (05x)",   r"\b05[0-9][-\s]?\d{7}\b",                    0.80),
        Pattern("IL Mobile (07x)",   r"\b07[2-9][-\s]?\d{7}\b",                    0.80),
        Pattern("IL Landline (0x)",  r"\b0[2-9][-\s]?\d{7}\b",                     0.65),
    ]
    CONTEXT = [
        "phone", "tel", "mobile", "cell", "fax", "telephone",
        "\u05d8\u05dc\u05e4\u05d5\u05df", "\u05e0\u05d9\u05d9\u05d3",
        "\u05e4\u05e7\u05e1", "\u05de\u05e1\u05e4\u05e8 \u05d8\u05dc\u05e4\u05d5\u05df",
        "\u05e1\u05dc\u05d5\u05dc\u05e8\u05d9",
    ]

    def __init__(self, lang="he"):
        super().__init__(
            supported_entity="IL_PHONE",
            patterns=self.PATTERNS,
            context=self.CONTEXT,
            supported_language=lang,
        )


class HebrewDateRecognizer(PatternRecognizer):
    PATTERNS = [
        Pattern("HE Date (slash)",  r"\b\d{1,2}/\d{1,2}/\d{4}\b",  0.75),
        Pattern("HE Date (dot)",    r"\b\d{1,2}\.\d{1,2}\.\d{4}\b", 0.75),
    ]
    CONTEXT = [
        "\u05ea\u05d0\u05e8\u05d9\u05da", "\u05e0\u05d5\u05dc\u05d3",
        "\u05dc\u05d9\u05d3\u05d4", "date", "born", "birth",
        "\u05e0\u05d7\u05ea\u05dd", "\u05de\u05d9\u05d5\u05dd", "\u05d1\u05d9\u05d5\u05dd",
        "\u05ea\u05d0\u05e8\u05d9\u05da \u05dc\u05d9\u05d3\u05d4",
    ]

    def __init__(self, lang="he"):
        super().__init__(
            supported_entity="DATE_TIME",
            patterns=self.PATTERNS,
            context=self.CONTEXT,
            supported_language=lang,
        )


class HebrewNameRecognizer(EntityRecognizer):
    """Dictionary-based recognizer for Hebrew-script person names."""
    SUPPORTED_ENTITY = "HE_PERSON"

    def __init__(self):
        super().__init__(
            supported_entities=[self.SUPPORTED_ENTITY],
            supported_language="he",
            name="HebrewNameRecognizer",
        )
        self._first = HEBREW_FIRST_NAMES
        self._last  = HEBREW_LAST_NAMES
        self._all   = HEBREW_FIRST_NAMES | HEBREW_LAST_NAMES

    def load(self):
        pass

    def analyze(self, text: str, entities: List[str], nlp_artifacts=None) -> List[RecognizerResult]:
        results = []
        tokens = re.finditer(r'[\u05d0-\u05ea\u05f0-\u05f4\ufb1d-\ufb4e]+', text)
        token_list = [(m.group(), m.start(), m.end()) for m in tokens]
        i = 0
        while i < len(token_list):
            word, start, end = token_list[i]
            if i + 1 < len(token_list):
                next_word, next_start, next_end = token_list[i + 1]
                if (word in self._first and next_word in self._last) or \
                   (word in self._last  and next_word in self._first) or \
                   (word in self._first and next_word in self._first):
                    results.append(RecognizerResult(
                        entity_type=self.SUPPORTED_ENTITY,
                        start=start, end=next_end, score=0.82,
                    ))
                    i += 2
                    continue
            if word in self._all:
                score = 0.75 if word in self._first else 0.70
                results.append(RecognizerResult(
                    entity_type=self.SUPPORTED_ENTITY,
                    start=start, end=end, score=score,
                ))
            i += 1
        return results


class HebrewLocationRecognizer(EntityRecognizer):
    """Dictionary-based recognizer for Hebrew-script locations."""
    SUPPORTED_ENTITY = "HE_LOCATION"

    def __init__(self):
        super().__init__(
            supported_entities=[self.SUPPORTED_ENTITY],
            supported_language="he",
            name="HebrewLocationRecognizer",
        )
        self._locations = sorted(HEBREW_LOCATIONS, key=len, reverse=True)

    def load(self):
        pass

    def analyze(self, text: str, entities: List[str], nlp_artifacts=None) -> List[RecognizerResult]:
        results = []
        for loc in self._locations:
            for m in re.finditer(re.escape(loc), text):
                results.append(RecognizerResult(
                    entity_type=self.SUPPORTED_ENTITY,
                    start=m.start(), end=m.end(), score=0.78,
                ))
        return results


# ---------------------------------------------------------------------------
#  File I/O helpers
# ---------------------------------------------------------------------------

def read_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        for enc in ("utf-8", "utf-8-sig", "windows-1255", "iso-8859-8", "cp1255"):
            try:
                with open(path, "r", encoding=enc) as fh:
                    return fh.read()
            except (UnicodeDecodeError, LookupError):
                continue
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()
    elif ext == ".docx":
        if not DOCX_OK:
            raise RuntimeError("python-docx not installed.")
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        if not PDF_OK:
            raise RuntimeError("pdfminer.six not installed.")
        return pdf_extract_text(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def write_file(path: str, text: str) -> None:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        if not DOCX_OK:
            raise RuntimeError("python-docx not installed.")
        doc = DocxDocument()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(path)
    else:
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(text)


# ---------------------------------------------------------------------------
#  PII engine
# ---------------------------------------------------------------------------

class PIIEngine:
    _instance = None

    @classmethod
    def get(cls):
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def __init__(self):
        self._en_analyzer = self._build_en_analyzer()
        self._he_analyzer = self._build_he_analyzer()
        self._xx_available = False

    def _build_en_analyzer(self) -> AnalyzerEngine:
        provider = NlpEngineProvider(nlp_configuration={
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
        })
        nlp_engine = provider.create_engine()
        registry = RecognizerRegistry()
        registry.load_predefined_recognizers(nlp_engine=nlp_engine)
        registry.add_recognizer(IsraeliIdRecognizer(lang="en"))
        registry.add_recognizer(IsraeliPhoneRecognizer(lang="en"))
        registry.add_recognizer(HebrewDateRecognizer(lang="en"))
        return AnalyzerEngine(nlp_engine=nlp_engine, registry=registry, supported_languages=["en"])

    def _build_he_analyzer(self) -> AnalyzerEngine:
        """Build the Hebrew analyzer engine.

        Strategy: always build a registry where every recognizer is registered
        under the SAME language code as the engine's supported_languages list.
        Never call load_predefined_recognizers() because it registers everything
        as 'en', which causes a Presidio validation error when the engine is
        declared as 'he'.
        """
        from presidio_analyzer.predefined_recognizers import (
            EmailRecognizer, IbanRecognizer, CreditCardRecognizer,
            UrlRecognizer, IpRecognizer,
        )
        try:
            from presidio_analyzer.predefined_recognizers import PhoneRecognizer
            _has_phone = True
        except Exception:
            _has_phone = False

        # --- Try to load the multilingual xx model (bundled in the .exe) ---
        nlp_engine = None
        try:
            import spacy
            spacy.load("xx_ent_wiki_sm")
            provider = NlpEngineProvider(nlp_configuration={
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "he", "model_name": "xx_ent_wiki_sm"}],
            })
            nlp_engine = provider.create_engine()
            self._xx_available = True
        except Exception:
            self._xx_available = False

        # Choose the language tag that will be used for EVERY recognizer
        # and for the AnalyzerEngine — they MUST match.
        if nlp_engine:
            lang = "he"
        else:
            # Fall back to en_core_web_sm; use "en" throughout
            lang = "en"
            try:
                provider2 = NlpEngineProvider(nlp_configuration={
                    "nlp_engine_name": "spacy",
                    "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
                })
                nlp_engine = provider2.create_engine()
            except Exception:
                return self._en_analyzer  # absolute last resort

        # Build a clean registry — all recognizers registered under `lang`
        registry = RecognizerRegistry()

        # Custom Israeli / Hebrew recognizers
        registry.add_recognizer(IsraeliIdRecognizer(lang=lang))
        registry.add_recognizer(IsraeliPhoneRecognizer(lang=lang))
        registry.add_recognizer(HebrewDateRecognizer(lang=lang))

        # Dictionary-based Hebrew NER (only useful when lang=="he", but harmless otherwise)
        if lang == "he":
            registry.add_recognizer(HebrewNameRecognizer())
            registry.add_recognizer(HebrewLocationRecognizer())

        # Standard Presidio recognizers — all registered under `lang`
        for rec_cls in (EmailRecognizer, IbanRecognizer, CreditCardRecognizer,
                        UrlRecognizer, IpRecognizer):
            try:
                registry.add_recognizer(rec_cls(supported_language=lang))
            except Exception:
                pass

        if _has_phone:
            try:
                registry.add_recognizer(
                    PhoneRecognizer(supported_language=lang, supported_regions=["IL"])
                )
            except Exception:
                try:
                    registry.add_recognizer(PhoneRecognizer(supported_language=lang))
                except Exception:
                    pass

        # NLP-backed recognizers for PERSON / LOCATION / ORG via spaCy NER
        from presidio_analyzer.predefined_recognizers import SpacyRecognizer
        try:
            registry.add_recognizer(
                SpacyRecognizer(
                    supported_language=lang,
                    supported_entities=["PERSON", "LOCATION", "ORGANIZATION",
                                        "DATE_TIME", "NRP"],
                    check_label_groups=[
                        ({"PERSON"},   {"PER", "PERSON"}),
                        ({"LOCATION"}, {"LOC", "GPE", "FAC"}),
                        ({"ORGANIZATION"}, {"ORG"}),
                        ({"DATE_TIME"}, {"DATE", "TIME"}),
                        ({"NRP"}, {"NORP"}),
                    ],
                )
            )
        except Exception:
            pass

        return AnalyzerEngine(
            nlp_engine=nlp_engine,
            registry=registry,
            supported_languages=[lang],
        )

    @staticmethod
    def _map_entity(entity_type: str) -> str:
        mapping = {
            "PERS": "PERSON", "PER": "PERSON",
            "LOC": "LOCATION", "GPE": "LOCATION",
            "ORG": "ORGANIZATION",
            "DATE": "DATE_TIME", "TIME": "DATE_TIME",
            "HE_PERSON": "PERSON",
            "HE_LOCATION": "LOCATION",
        }
        return mapping.get(entity_type, entity_type)

    @staticmethod
    def _resolve_overlaps(results):
        sorted_r = sorted(results, key=lambda r: (r.score, r.end - r.start), reverse=True)
        kept = []
        for cand in sorted_r:
            if not any(cand.start < a.end and cand.end > a.start for a in kept):
                kept.append(cand)
        kept.sort(key=lambda r: r.start)
        return kept

    def anonymize(self, text: str, confidence: float = DEFAULT_CONFIDENCE,
                  entities: Optional[List[str]] = None) -> Tuple[str, Dict, List]:
        if entities is None:
            entities = ALL_ENTITIES

        lang = detect_language(text)
        char_count = len(text)
        logger.info("Anonymize: lang=%s, chars=%d, confidence=%.2f, entities=%d",
                    lang, char_count, confidence, len(entities))

        if lang == "he":
            try:
                he_results = self._he_analyzer.analyze(text=text, language="he", score_threshold=confidence)
                logger.debug("Hebrew analyzer: %d raw results", len(he_results))
            except Exception as exc:
                logger.warning("Hebrew analyzer failed: %s", exc)
                he_results = []
            try:
                en_results = self._en_analyzer.analyze(text=text, language="en", score_threshold=confidence)
                logger.debug("English analyzer (Hebrew doc): %d raw results", len(en_results))
            except Exception as exc:
                logger.warning("English analyzer (Hebrew doc) failed: %s", exc)
                en_results = []
            raw = list(he_results) + list(en_results)
            analysis_lang = "he"
        else:
            try:
                raw = self._en_analyzer.analyze(text=text, language="en", entities=entities, score_threshold=confidence)
            except Exception as exc:
                logger.warning("English analyzer with entity filter failed (%s), retrying without filter", exc)
                raw = self._en_analyzer.analyze(text=text, language="en", score_threshold=confidence)
            analysis_lang = "en"
            logger.debug("English analyzer: %d raw results", len(raw))

        resolved = self._resolve_overlaps(raw)
        logger.debug("After overlap resolution: %d results", len(resolved))
        mapping: Dict[str, str] = {}
        entity_counts: Dict[str, int] = {}
        value_to_ph: Dict[str, str] = {}
        detections: List[dict] = []

        for result in sorted(resolved, key=lambda r: r.start, reverse=True):
            original = text[result.start:result.end]
            etype = self._map_entity(result.entity_type)
            label = ENTITY_LABELS.get(etype, etype)
            if original in value_to_ph:
                ph = value_to_ph[original]
            else:
                entity_counts[label] = entity_counts.get(label, 0) + 1
                ph = f"{{{{{label}_{entity_counts[label]}}}}}"
                mapping[ph] = original
                value_to_ph[original] = ph
            # Privacy-safe debug log: placeholder and label only, NOT the original value
            logger.debug("  Detected: %s  type=%-15s  score=%.2f  lang=%s",
                         ph, label, result.score, analysis_lang)
            detections.append({
                "placeholder": ph, "original": original,
                "label": label, "score": result.score,
                "start": result.start, "end": result.end,
                "lang": analysis_lang,
            })
            text = text[:result.start] + ph + text[result.end:]

        logger.info("Anonymize complete: %d unique PII items, %d total detections",
                    len(mapping), len(detections))
        return text, mapping, detections

    @staticmethod
    def restore(text: str, mapping: Dict[str, str]) -> str:
        logger.info("Restore: %d placeholders to replace", len(mapping))
        for ph, orig in mapping.items():
            text = text.replace(ph, orig)
        logger.info("Restore complete")
        return text

    @property
    def hebrew_ner_model(self) -> str:
        if self._xx_available:
            return "xx_ent_wiki_sm + dictionary (700+ names)"
        return "dictionary (700+ Israeli names)"


# ---------------------------------------------------------------------------
#  Custom PII project file helpers
# ---------------------------------------------------------------------------

CUSTOM_PII_FILENAME = "_custom_pii.json"

def get_project_pii_path(folder: str) -> str:
    """Return the path to the custom PII file for a given project folder."""
    return os.path.join(folder, CUSTOM_PII_FILENAME)

def load_project_pii(folder: str) -> List[dict]:
    """
    Load the project-level custom PII list from the folder.
    Returns a list of dicts: [{"text": str, "label": str}, ...]
    """
    path = get_project_pii_path(folder)
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            if isinstance(data, list):
                return data
        except Exception:
            pass
    return []

def save_project_pii(folder: str, entries: List[dict]) -> None:
    """Save the project-level custom PII list to the folder."""
    path = get_project_pii_path(folder)
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(entries, fh, indent=2, ensure_ascii=False)

def apply_custom_pii(text: str, entries: List[dict],
                     mapping: Dict[str, str],
                     entity_counts: Dict[str, int],
                     value_to_ph: Dict[str, str],
                     detections: List[dict]) -> str:
    """
    Apply manual custom PII entries to the text (case-sensitive exact match).
    Modifies mapping, entity_counts, value_to_ph, and detections in-place.
    Returns the modified text.
    """
    for entry in entries:
        raw = entry.get("text", "").strip()
        label = entry.get("label", "CUSTOM").upper().replace(" ", "_")
        if not raw:
            continue
        # Find all occurrences (longest first to avoid partial overlaps)
        for m in sorted(re.finditer(re.escape(raw), text),
                        key=lambda x: x.start(), reverse=True):
            original = m.group()
            if original in value_to_ph:
                ph = value_to_ph[original]
            else:
                entity_counts[label] = entity_counts.get(label, 0) + 1
                ph = f"{{{{{label}_{entity_counts[label]}}}}}"
                mapping[ph] = original
                value_to_ph[original] = ph
            detections.append({
                "placeholder": ph, "original": original,
                "label": label, "score": 1.0,
                "start": m.start(), "end": m.end(),
                "lang": "custom",
            })
            text = text[:m.start()] + ph + text[m.end():]
    return text


# ---------------------------------------------------------------------------
#  Translation map helpers
# ---------------------------------------------------------------------------

TRANSLATION_MAP_FILENAME = "_mapping_translated.json"


def _get_translation_map_path(mapping_path: str) -> str:
    """Return the sibling translation-map path for a given mapping.json path."""
    base = os.path.splitext(mapping_path)[0]
    # Strip trailing '_mapping' if present, then add '_mapping_translated'
    if base.endswith("_mapping"):
        base = base[:-len("_mapping")]
    return base + "_mapping_translated.json"


def generate_translation_map(mapping: Dict[str, str]) -> Dict[str, dict]:
    """
    Build a translation map from an anonymization mapping.

    The translation map has the structure:
        { placeholder: { "original": str, "translated": str, "entity": str } }

    The 'translated' field is pre-filled with:
      - A known English transliteration for Hebrew names/locations (from HEBREW_TRANSLITERATIONS)
      - The original value unchanged for structured PII (IDs, phones, emails, dates, IBANs)
      - An empty string for anything else (user must fill in manually)
    """
    try:
        from hebrew_data import HEBREW_TRANSLITERATIONS
    except ImportError:
        HEBREW_TRANSLITERATIONS = {}

    # Entity types where the value should be copied unchanged (same in both languages)
    COPY_AS_IS = {"IL_ID", "PHONE_NUMBER", "PHONE", "EMAIL_ADDRESS", "EMAIL",
                  "IBAN_CODE", "IBAN", "CREDIT_CARD", "DATE_TIME", "DATE",
                  "IL_PHONE", "IL_IBAN", "IL_DATE", "US_SSN", "US_PASSPORT"}

    result = {}
    for ph, original in mapping.items():
        # Extract entity type from placeholder like {{PERSON_1}}
        m = re.match(r'\{\{([A-Z_]+)_(\d+)\}\}', ph)
        entity = m.group(1) if m else "CUSTOM"

        # Determine pre-filled translation
        if entity in COPY_AS_IS:
            # Structured PII — keep the same value in both languages
            translated = original
        else:
            # Try to look up word by word in the transliteration dict
            # This handles multi-word names like "דוד לוי" -> "David Levy"
            words = original.split()
            translated_words = []
            all_found = True
            for word in words:
                clean = word.strip(".,;:()\"'")
                if clean in HEBREW_TRANSLITERATIONS:
                    translated_words.append(HEBREW_TRANSLITERATIONS[clean])
                else:
                    all_found = False
                    translated_words.append("")
            if all_found and translated_words:
                translated = " ".join(translated_words)
            elif any(t for t in translated_words):
                # Partial match — join what we have, blanks for unknowns
                translated = " ".join(t if t else "?" for t in translated_words)
            else:
                translated = ""

        result[ph] = {
            "original": original,
            "translated": translated,
            "entity": entity,
        }
    return result


# ---------------------------------------------------------------------------
#  Reusable styled widgets
# ---------------------------------------------------------------------------

def styled_button(parent, text, command, bg=ACCENT, fg="white", width=None, pady=8):
    kw = dict(text=text, command=command, bg=bg, fg=fg,
              font=("Segoe UI", 10, "bold"), relief="flat",
              activebackground=ACCENT_HOVER, activeforeground="white",
              cursor="hand2", pady=pady, bd=0)
    if width:
        kw["width"] = width
    btn = tk.Button(parent, **kw)
    btn.bind("<Enter>", lambda e: btn.config(bg=ACCENT_HOVER))
    btn.bind("<Leave>", lambda e: btn.config(bg=bg))
    return btn


def section_label(parent, text):
    return tk.Label(parent, text=text, bg=PANEL_BG, fg=ACCENT, font=("Segoe UI", 9, "bold"))


def entry_row(parent, label_text, var, browse_cmd=None):
    row = tk.Frame(parent, bg=PANEL_BG)
    tk.Label(row, text=label_text, bg=PANEL_BG, fg=TEXT_MAIN,
             font=("Segoe UI", 9), width=18, anchor="w").pack(side="left")
    ent = tk.Entry(row, textvariable=var, bg=ENTRY_BG, fg=TEXT_MAIN,
                   insertbackground=TEXT_MAIN, relief="flat", font=("Segoe UI", 9), bd=4)
    ent.pack(side="left", fill="x", expand=True)
    if browse_cmd:
        tk.Button(row, text="Browse", command=browse_cmd, bg=BORDER, fg=TEXT_MAIN,
                  font=("Segoe UI", 8), relief="flat", cursor="hand2", padx=6
                  ).pack(side="left", padx=(4, 0))
    return row


# ---------------------------------------------------------------------------
#  Splash screen
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
#  Auto-Updater
# ---------------------------------------------------------------------------

class Updater:
    """
    Checks GitHub Releases for a newer version of the app.
    Runs entirely in background threads so the UI is never blocked.

    Flow:
      1. On startup, check_in_background() fetches version.json from the
         latest GitHub Release.
      2. If a newer version is found, it calls on_update_available(info)
         on the main thread via tk.after().
      3. The App shows a non-intrusive banner with a "Download & Restart" button.
      4. download_and_restart() streams the new .exe, replaces the current
         executable, and relaunches the process.
    """

    def __init__(self, app: "App"):
        self._app = app
        self._banner_frame: Optional[tk.Frame] = None

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def check_in_background(self):
        """Spawn a daemon thread to check for updates silently."""
        threading.Thread(target=self._check, daemon=True).start()

    def download_and_restart(self, download_url: str, new_version: str):
        """Show a progress dialog and replace the running .exe."""
        dlg = _UpdateDownloadDialog(self._app, new_version)
        threading.Thread(
            target=self._do_download,
            args=(download_url, new_version, dlg),
            daemon=True,
        ).start()

    # ------------------------------------------------------------------
    # Internal
    # ------------------------------------------------------------------

    def _check(self):
        try:
            req = urllib.request.Request(
                UPDATE_MANIFEST_URL,
                headers={"User-Agent": f"PII-Processor/{APP_VERSION}"},
            )
            ctx = _make_ssl_context()
            with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
                data = json.loads(resp.read().decode())
            latest = data.get("version", "")
            url    = data.get("download_url", "")
            notes  = data.get("release_notes", "")
            logger.debug("Update check: latest=%s current=%s", latest, APP_VERSION)
            if latest and url and self._is_newer(latest):
                self._app.after(0, lambda: self._on_update_available(
                    latest, url, notes
                ))
        except Exception as exc:
            logger.warning("Background update check failed: %s: %s",
                           type(exc).__name__, exc)  # logged but not shown to user

    @staticmethod
    def _is_newer(remote: str) -> bool:
        """Return True if remote version string is greater than APP_VERSION."""
        def _parts(v):
            try:
                return tuple(int(x) for x in str(v).split("."))
            except Exception:
                return (0,)
        return _parts(remote) > _parts(APP_VERSION)

    def _on_update_available(self, version: str, url: str, notes: str):
        """Show the update banner in the app header."""
        if self._banner_frame:
            return  # already shown
        hdr = self._app._header_frame
        self._banner_frame = tk.Frame(hdr, bg="#2a6a2a", padx=8, pady=4)
        self._banner_frame.pack(side="right", padx=(0, 8))

        tk.Label(
            self._banner_frame,
            text=f"\u2b06  Update available  v{version}",
            bg="#2a6a2a", fg="#90ee90",
            font=("Segoe UI", 9, "bold"),
        ).pack(side="left", padx=(0, 8))

        tk.Button(
            self._banner_frame,
            text="Download & Restart",
            bg="#3cb371", fg="white",
            font=("Segoe UI", 9, "bold"),
            relief="flat", cursor="hand2", padx=8,
            command=lambda: self.download_and_restart(url, version),
        ).pack(side="left")

        if notes:
            tk.Label(
                self._banner_frame,
                text=f"({notes})",
                bg="#2a6a2a", fg="#90ee90",
                font=("Segoe UI", 8),
            ).pack(side="left", padx=(6, 0))

    def _do_download(self, url: str, version: str, dlg: "_UpdateDownloadDialog"):
        """Download the new .exe and replace the current one, then restart."""
        try:
            # Determine the path of the running executable
            if getattr(sys, "frozen", False):
                current_exe = sys.executable
            else:
                current_exe = os.path.abspath(sys.argv[0])

            # Download to a temp file next to the current exe
            exe_dir = os.path.dirname(current_exe)
            tmp_path = os.path.join(exe_dir, f"_update_{version}.exe")

            ctx = _make_ssl_context()
            req = urllib.request.Request(
                url,
                headers={"User-Agent": f"PII-Processor/{APP_VERSION}"},
            )
            with urllib.request.urlopen(req, timeout=120, context=ctx) as resp:
                total = int(resp.headers.get("Content-Length", 0))
                downloaded = 0
                chunk = 65536
                with open(tmp_path, "wb") as f:
                    while True:
                        buf = resp.read(chunk)
                        if not buf:
                            break
                        f.write(buf)
                        downloaded += len(buf)
                        if total > 0:
                            pct = int(downloaded * 100 / total)
                            self._app.after(0, lambda p=pct: dlg.set_progress(p))

            # Rename current exe to .bak, rename new exe to current name
            backup = current_exe + ".bak"
            if os.path.exists(backup):
                os.remove(backup)
            os.rename(current_exe, backup)
            os.rename(tmp_path, current_exe)

            # Relaunch and exit
            self._app.after(0, lambda: dlg.set_status("Restarting..."))
            self._app.after(500, lambda: self._restart(current_exe))

        except urllib.error.URLError as exc:
            reason = exc.reason if exc.reason else str(exc)
            logger.error("Update download URLError: %s", reason)
            self._app.after(0, lambda: dlg.destroy())
            self._app.after(
                0,
                lambda r=str(reason): messagebox.showerror(
                    "Update Failed",
                    f"Network error while downloading update:\n{r}\n\n"
                    "Please check your internet connection and try again,\n"
                    "or download the latest version manually from GitHub.",
                ),
            )
        except Exception as exc:
            err = f"{type(exc).__name__}: {exc}"
            logger.error("Update download failed: %s", err)
            self._app.after(0, lambda: dlg.destroy())
            self._app.after(
                0,
                lambda e=err: messagebox.showerror(
                    "Update Failed",
                    f"Could not download update:\n{e}\n\n"
                    "Please download the latest version manually from GitHub.",
                ),
            )

    @staticmethod
    def _restart(exe_path: str):
        subprocess.Popen([exe_path])
        sys.exit(0)


class _UpdateDownloadDialog(tk.Toplevel):
    """Modal-ish progress window shown while downloading an update."""

    def __init__(self, parent, version: str):
        super().__init__(parent)
        self.title("Downloading Update")
        self.geometry("400x140")
        self.resizable(False, False)
        self.configure(bg="#1e1e2e")
        self.grab_set()

        tk.Label(
            self,
            text=f"Downloading PII Processor v{version}...",
            bg="#1e1e2e", fg="#cdd6f4",
            font=("Segoe UI", 11, "bold"),
        ).pack(pady=(24, 8))

        self._bar = ttk.Progressbar(self, length=340, maximum=100)
        self._bar.pack(pady=4)

        self._lbl = tk.Label(
            self, text="0%",
            bg="#1e1e2e", fg="#a6adc8",
            font=("Segoe UI", 9),
        )
        self._lbl.pack()

    def set_progress(self, pct: int):
        self._bar["value"] = pct
        self._lbl.config(text=f"{pct}%")
        self.update_idletasks()

    def set_status(self, msg: str):
        self._lbl.config(text=msg)
        self.update_idletasks()


# ---------------------------------------------------------------------------
#  Hebrew ambiguity detection
# ---------------------------------------------------------------------------

# Hebrew title/honorific words that strongly suggest the following word is a name
HEBREW_TITLE_WORDS = {
    "אדון", "גברת", "גב", "רב", "הרב", "שופט", "שופטת", "עורך", "עורכת",
    "מר", "הנאשם", "הנאשמת", "התובע", "התובעת", "הנתבע", "הנתבעת",
    "המבקש", "המבקשת", "המשיב", "המשיבה", "הצד", "הלקוח", "הלקוחה",
    "פרופסור", "דוקטור", "עורך דין", "עורכת דין",
    # With geresh/gershayim (unicode escapes to avoid string literal issues)
    "\u05d2\u05d1\u05f3",   # גב׳
    "\u05d3\u05f4\u05e8",   # ד״ר
    "\u05e2\u05d5\u05f4\u05d3",  # עו״ד
    "\u05e2\u05d5\u05d4\u05f4\u05d3",  # עוה״ד
    "\u05e4\u05e8\u05d5\u05e4\u05f3",  # פרופ׳
    "\u05d3\u05e8\u05f3",   # דר׳
}

# Hebrew words that are BOTH common nouns AND names — the ambiguous set
# These are words from the dictionary that are also everyday Hebrew words
HEBREW_AMBIGUOUS_NAMES = {
    # Names that are also nouns/adjectives
    "אביב", "אביגיל", "אור", "אורה", "אורי", "אורן", "אורית",
    "אלה", "אלון", "אמיר", "אסף", "ארז", "ארי", "אריאל", "אריה",
    "בועז", "בן", "בר", "גל", "גלי", "גפן", "דביר", "דגן", "דוד",
    "דור", "דנה", "דפנה", "הדס", "הדסה", "הילה", "זהר", "זיו",
    "חן", "חנה", "טל", "יאיר", "יובל", "יונה", "יונתן", "יורם",
    "יחיאל", "ים", "יניב", "יעל", "יצחק", "ירדן", "כרמל", "לב",
    "לי", "לילה", "לימור", "מור", "מיה", "מיכל", "מירב", "מירי",
    "מנחם", "מעיין", "מרב", "נדב", "נוי", "נועה", "נועם", "נורית",
    "נטע", "ניב", "ניר", "נעם", "נעמי", "נתן", "סהר", "סיון",
    "עדי", "עומר", "עידן", "עינב", "עמית", "עמנואל", "ענבל",
    "ענת", "עפר", "ערן", "פנינה", "צבי", "צור", "קרן", "רוי",
    "רון", "רונה", "רונית", "רות", "רז", "ריבה", "רם", "שגיא",
    "שחר", "שי", "שיר", "שירה", "שלג", "שלום", "שמש", "שני",
    "שקד", "תהל", "תום", "תמר",
    # Locations that are also common words
    "רמת", "גן", "עמק", "מרכז", "צפון", "דרום", "מזרח", "מערב",
    "שרון", "כרמל", "גלבוע", "עמק",
}


def find_hebrew_ambiguous_candidates(
    original_text: str,
    already_mapped_values: set,
) -> List[dict]:
    """
    Scan a Hebrew document for tokens that:
      1. Appear in the Hebrew name/location dictionaries, AND
      2. Were NOT already detected by the NLP engine, AND
      3. Are ambiguous (also common Hebrew nouns)

    Returns a list of candidate dicts:
      {"text": str, "label": str, "context": str, "count": int}
    sorted by frequency descending.
    """
    if not HEBREW_FIRST_NAMES and not HEBREW_LAST_NAMES and not HEBREW_LOCATIONS:
        return []

    all_names = HEBREW_FIRST_NAMES | HEBREW_LAST_NAMES
    candidates: Dict[str, dict] = {}

    # Split into sentences for context extraction
    sentences = re.split(r'[.!?\n]+', original_text)
    sentence_map: Dict[int, str] = {}  # char_offset -> sentence text
    pos = 0
    for sent in sentences:
        for i in range(pos, pos + len(sent) + 1):
            sentence_map[i] = sent.strip()
        pos += len(sent) + 1

    # Tokenize: split on whitespace and punctuation, keep Hebrew words
    hebrew_word_re = re.compile(r'[\u05d0-\u05ea\ufb1d-\ufb4e\'\"״׳]+')

    for m in hebrew_word_re.finditer(original_text):
        word = m.group().strip("'\"״׳")
        if len(word) < 2:
            continue
        # Skip if already replaced by NLP engine
        if word in already_mapped_values:
            continue

        suggested_label = None
        is_ambiguous = word in HEBREW_AMBIGUOUS_NAMES

        # Check if it's a name (first or last)
        if word in all_names:
            suggested_label = "PERSON"
        elif word in HEBREW_LOCATIONS:
            suggested_label = "LOCATION"

        if suggested_label is None:
            continue

        # Check if preceded by a title word (strong signal → include even if not ambiguous)
        start = m.start()
        preceding = original_text[max(0, start - 20):start].strip()
        preceding_words = preceding.split()
        preceded_by_title = any(w.strip("'\"״׳") in HEBREW_TITLE_WORDS
                                 for w in preceding_words[-3:])

        # Include if: (a) preceded by title, OR (b) ambiguous word in dictionary
        if not preceded_by_title and not is_ambiguous:
            continue

        # Get context sentence
        context = sentence_map.get(start, "")[:120]
        if not context:
            context = original_text[max(0, start - 40):start + len(word) + 40]

        if word not in candidates:
            candidates[word] = {
                "text": word,
                "label": suggested_label,
                "context": context,
                "count": 0,
                "preceded_by_title": preceded_by_title,
            }
        candidates[word]["count"] += 1
        if preceded_by_title:
            candidates[word]["preceded_by_title"] = True

    # Sort: title-preceded first, then by frequency
    result = sorted(
        candidates.values(),
        key=lambda c: (not c["preceded_by_title"], -c["count"])
    )
    logger.debug("Hebrew ambiguity scan: %d candidates found", len(result))
    return result


class HebrewReviewDialog(tk.Toplevel):
    """
    Modal dialog shown after Hebrew NLP detection.
    Presents a checklist of ambiguous Hebrew words so the user can
    confirm which ones are PII (names/locations) before anonymization.
    """

    def __init__(self, parent, candidates: List[dict]):
        super().__init__(parent)
        self.title("Hebrew PII Review — Confirm Ambiguous Terms")
        self.configure(bg=DARK_BG)
        self.resizable(True, True)
        self.grab_set()  # modal

        # Size and centre
        w, h = 780, 560
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        self.geometry(f"{w}x{h}+{(sw - w) // 2}+{(sh - h) // 2}")
        self.minsize(600, 400)

        self._candidates = candidates
        self._vars: List[tk.BooleanVar] = []
        self._label_vars: List[tk.StringVar] = []
        self.approved: List[dict] = []  # filled on Confirm
        self._confirmed = False

        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_skip)

    def _build_ui(self):
        # ── Header ──────────────────────────────────────────────────────────
        hdr = tk.Frame(self, bg=PANEL_BG)
        hdr.pack(fill="x")
        tk.Label(
            hdr,
            text="Hebrew PII Review",
            bg=PANEL_BG, fg=ACCENT,
            font=("Segoe UI", 13, "bold"),
        ).pack(side="left", padx=16, pady=10)

        tk.Label(
            hdr,
            text=f"{len(self._candidates)} ambiguous term(s) found",
            bg=PANEL_BG, fg=TEXT_DIM,
            font=("Segoe UI", 9),
        ).pack(side="right", padx=16)

        # ── Instruction ──────────────────────────────────────────────────────
        tk.Label(
            self,
            text=(
                "The following Hebrew words were found in the document. They appear in the name/location "
                "dictionary but may also be common nouns.\n"
                "Tick the checkbox next to each word that IS a person name or location that should be anonymized."
            ),
            bg=DARK_BG, fg=TEXT_MAIN,
            font=("Segoe UI", 9),
            wraplength=740, justify="left",
        ).pack(fill="x", padx=16, pady=(8, 4))

        # ── Select all / none ────────────────────────────────────────────────
        ctrl = tk.Frame(self, bg=DARK_BG)
        ctrl.pack(fill="x", padx=16, pady=(0, 4))
        tk.Button(ctrl, text="Select All", bg=PANEL_BG, fg=TEXT_MAIN,
                  relief="flat", font=("Segoe UI", 9), cursor="hand2",
                  command=self._select_all).pack(side="left", padx=(0, 8))
        tk.Button(ctrl, text="Select None", bg=PANEL_BG, fg=TEXT_MAIN,
                  relief="flat", font=("Segoe UI", 9), cursor="hand2",
                  command=self._select_none).pack(side="left")

        # ── Scrollable candidate list ────────────────────────────────────────
        outer = tk.Frame(self, bg=DARK_BG)
        outer.pack(fill="both", expand=True, padx=16, pady=(0, 8))

        canvas = tk.Canvas(outer, bg=DARK_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        self._scroll_frame = tk.Frame(canvas, bg=DARK_BG)

        self._scroll_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=self._scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Bind mousewheel
        canvas.bind_all("<MouseWheel>",
            lambda e: canvas.yview_scroll(int(-1 * (e.delta / 120)), "units"))

        # Column headers
        hdr_row = tk.Frame(self._scroll_frame, bg=PANEL_BG)
        hdr_row.pack(fill="x", pady=(0, 2))
        tk.Label(hdr_row, text="Include", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Segoe UI", 8, "bold"), width=8).pack(side="left", padx=4)
        tk.Label(hdr_row, text="Hebrew Word", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Segoe UI", 8, "bold"), width=14).pack(side="left", padx=4)
        tk.Label(hdr_row, text="Count", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Segoe UI", 8, "bold"), width=6).pack(side="left", padx=4)
        tk.Label(hdr_row, text="Label", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Segoe UI", 8, "bold"), width=12).pack(side="left", padx=4)
        tk.Label(hdr_row, text="Context (surrounding text)", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Segoe UI", 8, "bold")).pack(side="left", padx=4)

        # Candidate rows
        LABEL_OPTIONS = ["PERSON", "LOCATION", "ORGANIZATION", "CUSTOM"]
        for i, cand in enumerate(self._candidates):
            row_bg = DARK_BG if i % 2 == 0 else PANEL_BG
            row = tk.Frame(self._scroll_frame, bg=row_bg)
            row.pack(fill="x", pady=1)

            # Pre-tick if preceded by title word
            default_checked = cand.get("preceded_by_title", False)
            var = tk.BooleanVar(value=default_checked)
            self._vars.append(var)

            tk.Checkbutton(
                row, variable=var,
                bg=row_bg, activebackground=row_bg,
                selectcolor=ENTRY_BG, fg=TEXT_MAIN,
                width=6,
            ).pack(side="left", padx=4)

            # Hebrew word (RTL display)
            tk.Label(
                row, text=cand["text"],
                bg=row_bg, fg=TAG_PERSON,
                font=("David", 11, "bold") if sys.platform == "win32" else ("Segoe UI", 11, "bold"),
                width=14, anchor="e",  # right-align for RTL
            ).pack(side="left", padx=4)

            # Count
            tk.Label(
                row, text=str(cand["count"]),
                bg=row_bg, fg=TEXT_DIM,
                font=("Segoe UI", 9), width=6,
            ).pack(side="left", padx=4)

            # Label dropdown
            lvar = tk.StringVar(value=cand["label"])
            self._label_vars.append(lvar)
            ttk.Combobox(
                row, textvariable=lvar,
                values=LABEL_OPTIONS, state="readonly",
                width=12, font=("Segoe UI", 9),
            ).pack(side="left", padx=4)

            # Context snippet
            ctx = cand.get("context", "")[:100]
            tk.Label(
                row, text=ctx,
                bg=row_bg, fg=TEXT_DIM,
                font=("Segoe UI", 8),
                anchor="w", justify="left",
                wraplength=340,
            ).pack(side="left", padx=4, fill="x", expand=True)

        # ── Footer buttons ───────────────────────────────────────────────────
        footer = tk.Frame(self, bg=PANEL_BG)
        footer.pack(fill="x", side="bottom")

        styled_button(
            footer, "✓  Confirm & Anonymize",
            command=self._on_confirm,
        ).pack(side="right", padx=12, pady=8)

        tk.Button(
            footer, text="Skip — Anonymize Without These",
            bg=PANEL_BG, fg=TEXT_DIM,
            relief="flat", font=("Segoe UI", 9),
            cursor="hand2",
            command=self._on_skip,
        ).pack(side="right", padx=4, pady=8)

        tk.Label(
            footer,
            text="Ticked items will be replaced with placeholders.",
            bg=PANEL_BG, fg=TEXT_DIM,
            font=("Segoe UI", 8),
        ).pack(side="left", padx=12)

    def _select_all(self):
        for v in self._vars:
            v.set(True)

    def _select_none(self):
        for v in self._vars:
            v.set(False)

    def _on_confirm(self):
        self.approved = [
            {"text": self._candidates[i]["text"],
             "label": self._label_vars[i].get()}
            for i, v in enumerate(self._vars) if v.get()
        ]
        logger.info("Hebrew review: user approved %d / %d candidates",
                    len(self.approved), len(self._candidates))
        self._confirmed = True
        self.grab_release()
        self.destroy()

    def _on_skip(self):
        self.approved = []
        logger.info("Hebrew review: user skipped (0 candidates approved)")
        self._confirmed = False
        self.grab_release()
        self.destroy()


class SplashScreen(tk.Toplevel):
    """Full-window splash shown while the NLP engine loads."""

    def __init__(self, master):
        super().__init__(master)
        self.overrideredirect(True)          # borderless
        self.configure(bg=DARK_BG)
        self.attributes("-topmost", True)

        # Centre on screen
        w, h = 520, 320
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        x = (sw - w) // 2
        y = (sh - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

        # Logo / title
        tk.Label(
            self, text="⚖️  Legal PII Anonymizer",
            bg=DARK_BG, fg=TEXT_MAIN,
            font=("Segoe UI", 20, "bold"),
        ).pack(pady=(48, 4))

        tk.Label(
            self, text="Hebrew \u2022 English  |  v3.0",
            bg=DARK_BG, fg=ACCENT,
            font=("Segoe UI", 11),
        ).pack(pady=(0, 32))

        # Animated progress bar
        self._bar = ttk.Progressbar(
            self, mode="indeterminate", length=380
        )
        self._bar.pack(pady=(0, 16))
        self._bar.start(12)

        # Status label
        self._lbl = tk.Label(
            self, text="Initialising NLP engine — please wait...",
            bg=DARK_BG, fg=TEXT_DIM,
            font=("Segoe UI", 9),
            wraplength=460,
        )
        self._lbl.pack()

        tk.Label(
            self,
            text="(First launch may take 10\u201330 seconds while Windows caches the application)",
            bg=DARK_BG, fg=TEXT_DIM,
            font=("Segoe UI", 8),
            wraplength=460,
        ).pack(pady=(8, 0))

    def set_status(self, msg: str):
        self._lbl.config(text=msg)

    def close(self):
        self._bar.stop()
        self.destroy()


# ---------------------------------------------------------------------------
#  Main Application Window
# ---------------------------------------------------------------------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(f"Legal Document PII Anonymizer  |  Hebrew / English  v{APP_VERSION}")
        self.geometry("1150x820")
        self.minsize(900, 640)
        self.configure(bg=DARK_BG)
        self._engine_ready = False
        self._updater = Updater(self)

        # Show splash immediately, hide main window until engine is ready
        self.withdraw()
        self._splash = SplashScreen(self)
        self.update()

        self._build_ui()
        self._load_engine_async()

    def _load_engine_async(self):
        self._set_status("Loading NLP engine...", WARNING)
        threading.Thread(target=self._load_engine, daemon=True).start()

    def _load_engine(self):
        try:
            logger.info("Loading NLP engine...")
            self.after(0, lambda: self._splash.set_status(
                "Loading English NLP model (en_core_web_sm)..."
            ))
            engine = PIIEngine.get()
            self._engine_ready = True
            model_info = engine.hebrew_ner_model
            msg = f"Engine ready — Hebrew: {model_info}"
            logger.info("Engine loaded successfully. Hebrew NER: %s", model_info)
            self.after(0, lambda: self._splash.set_status("Engine ready! Opening application..."))
            self.after(300, self._show_main)
            self.after(0, lambda: self._set_status(msg, SUCCESS))
            self.after(0, lambda: self._lang_lbl.config(text=f"EN + HE ({model_info})", fg=SUCCESS))
        except Exception as exc:
            logger.error("Engine load failed: %s", exc, exc_info=True)
            self.after(0, lambda: self._splash.set_status(f"Error: {exc}"))
            self.after(0, lambda: self._set_status(f"Engine error: {exc}", DANGER))
            self.after(1500, self._show_main)

    def _show_main(self):
        """Close splash, reveal the main window, then check for updates."""
        try:
            self._splash.close()
        except Exception:
            pass
        self.deiconify()
        self.lift()
        # Check for updates silently in the background after the window is shown
        self.after(2000, self._updater.check_in_background)

    def _build_ui(self):
        header = tk.Frame(self, bg=PANEL_BG, pady=14)
        header.pack(fill="x")
        self._header_frame = header  # Updater needs this reference
        tk.Label(header, text=f"Legal Document PII Anonymizer & Restorer  v{APP_VERSION}",
                 bg=PANEL_BG, fg=TEXT_MAIN, font=("Segoe UI", 14, "bold")).pack(side="left", padx=20)
        self._lang_lbl = tk.Label(header, text="Loading...", bg=PANEL_BG, fg=WARNING, font=("Segoe UI", 9))
        self._lang_lbl.pack(side="right", padx=20)
        self._status_lbl = tk.Label(header, text="", bg=PANEL_BG, fg=WARNING, font=("Segoe UI", 9))
        self._status_lbl.pack(side="right", padx=20)

        style = ttk.Style(self)
        style.theme_use("default")
        style.configure("TNotebook", background=DARK_BG, borderwidth=0)
        style.configure("TNotebook.Tab", background=PANEL_BG, foreground=TEXT_DIM,
                        font=("Segoe UI", 10, "bold"), padding=[18, 8])
        style.map("TNotebook.Tab", background=[("selected", ACCENT)], foreground=[("selected", "white")])

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=12, pady=(8, 0))

        self._tab_anon    = tk.Frame(nb, bg=DARK_BG)
        self._tab_restore = tk.Frame(nb, bg=DARK_BG)
        self._tab_transmap = tk.Frame(nb, bg=DARK_BG)
        self._tab_batch   = tk.Frame(nb, bg=DARK_BG)
        self._tab_custom  = tk.Frame(nb, bg=DARK_BG)
        self._tab_debug   = tk.Frame(nb, bg=DARK_BG)
        self._tab_about   = tk.Frame(nb, bg=DARK_BG)

        nb.add(self._tab_anon,     text="  \U0001f512  Anonymize  ")
        nb.add(self._tab_restore,  text="  \U0001f513  Restore  ")
        nb.add(self._tab_transmap, text="  \U0001f310  Translation Map  ")
        nb.add(self._tab_batch,    text="  \U0001f4c2  Batch Folder  ")
        nb.add(self._tab_custom,   text="  \u270f  Custom PII  ")
        nb.add(self._tab_debug,    text="  \U0001f41b  Debug Log  ")
        nb.add(self._tab_about,    text="  \u2139  About  ")

        self._build_anonymize_tab()
        self._build_restore_tab()
        self._build_translation_map_tab()
        self._build_batch_tab()
        self._build_custom_pii_tab()
        self._build_debug_tab()
        self._build_about_tab()

        bar = tk.Frame(self, bg=PANEL_BG, height=28)
        bar.pack(fill="x", side="bottom")
        self._progress = ttk.Progressbar(bar, mode="indeterminate", length=200)
        self._progress.pack(side="right", padx=12, pady=4)
        self._bar_lbl = tk.Label(bar, text="", bg=PANEL_BG, fg=TEXT_DIM, font=("Segoe UI", 8))
        self._bar_lbl.pack(side="left", padx=12)

    def _build_anonymize_tab(self):
        tab = self._tab_anon
        left = tk.Frame(tab, bg=DARK_BG, width=340)
        left.pack(side="left", fill="y", padx=(12, 6), pady=12)
        left.pack_propagate(False)
        right = tk.Frame(tab, bg=DARK_BG)
        right.pack(side="left", fill="both", expand=True, padx=(0, 12), pady=12)

        panel = tk.Frame(left, bg=PANEL_BG, bd=0)
        panel.pack(fill="both", expand=True, pady=(0, 8))

        section_label(panel, "INPUT / OUTPUT").pack(anchor="w", padx=14, pady=(14, 4))

        self._anon_input  = tk.StringVar()
        self._anon_output = tk.StringVar()
        self._anon_map    = tk.StringVar()

        for lbl, var, cmd in [
            ("Input document",    self._anon_input,  self._browse_input),
            ("Anonymized output", self._anon_output, self._browse_anon_out),
            ("Mapping file",      self._anon_map,    self._browse_anon_map),
        ]:
            entry_row(panel, lbl, var, cmd).pack(fill="x", padx=14, pady=3)

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)
        section_label(panel, "DETECTION SETTINGS").pack(anchor="w", padx=14, pady=(0, 6))

        lang_row = tk.Frame(panel, bg=PANEL_BG)
        lang_row.pack(fill="x", padx=14, pady=(0, 6))
        tk.Label(lang_row, text="Language", bg=PANEL_BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9), width=18, anchor="w").pack(side="left")
        self._lang_override = tk.StringVar(value="auto")
        ttk.Combobox(lang_row, textvariable=self._lang_override,
                     values=["auto", "English", "Hebrew"],
                     state="readonly", width=16, font=("Segoe UI", 9)).pack(side="left")

        conf_row = tk.Frame(panel, bg=PANEL_BG)
        conf_row.pack(fill="x", padx=14, pady=(0, 4))
        tk.Label(conf_row, text="Min. confidence", bg=PANEL_BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9), width=18, anchor="w").pack(side="left")
        self._confidence = tk.DoubleVar(value=DEFAULT_CONFIDENCE)
        tk.Scale(conf_row, from_=0.3, to=1.0, resolution=0.05, orient="horizontal",
                 variable=self._confidence, bg=PANEL_BG, fg=TEXT_MAIN, troughcolor=ENTRY_BG,
                 highlightthickness=0, font=("Segoe UI", 8), activebackground=ACCENT, length=120
                 ).pack(side="left")
        self._conf_lbl = tk.Label(conf_row, text=f"{DEFAULT_CONFIDENCE:.2f}",
                                   bg=PANEL_BG, fg=ACCENT, font=("Segoe UI", 9, "bold"), width=4)
        self._conf_lbl.pack(side="left")
        self._confidence.trace_add("write",
            lambda *_: self._conf_lbl.config(text=f"{self._confidence.get():.2f}"))

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)
        section_label(panel, "ENTITY TYPES TO DETECT").pack(anchor="w", padx=14, pady=(0, 6))

        self._entity_vars: Dict[str, tk.BooleanVar] = {}
        scroll_frame = tk.Frame(panel, bg=PANEL_BG)
        scroll_frame.pack(fill="x", padx=14)

        display_entities = {k: v for k, v in ENTITY_LABELS.items()
                            if k not in ("HE_PERSON", "HE_LOCATION")}
        for i, (etype, label) in enumerate(display_entities.items()):
            var = tk.BooleanVar(value=True)
            self._entity_vars[etype] = var
            col = i % 2
            row_idx = i // 2
            he_name = HE_ENTITY_NAMES.get(label, "")
            display = f"{label}" + (f" ({he_name})" if he_name else "")
            tk.Checkbutton(scroll_frame, text=display, variable=var,
                           bg=PANEL_BG, fg=TEXT_MAIN, selectcolor=ENTRY_BG,
                           activebackground=PANEL_BG, activeforeground=TEXT_MAIN,
                           font=("Segoe UI", 8), anchor="w"
                           ).grid(row=row_idx, column=col, sticky="w", pady=1)

        for k in ("HE_PERSON", "HE_LOCATION"):
            self._entity_vars[k] = tk.BooleanVar(value=True)

        sel_row = tk.Frame(panel, bg=PANEL_BG)
        sel_row.pack(fill="x", padx=14, pady=(6, 0))
        tk.Button(sel_row, text="Select All", command=self._select_all_entities,
                  bg=BORDER, fg=TEXT_MAIN, font=("Segoe UI", 8), relief="flat",
                  cursor="hand2", padx=6).pack(side="left", padx=(0, 4))
        tk.Button(sel_row, text="Clear All", command=self._clear_all_entities,
                  bg=BORDER, fg=TEXT_MAIN, font=("Segoe UI", 8), relief="flat",
                  cursor="hand2", padx=6).pack(side="left")

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)
        styled_button(panel, "\U0001f512  Remove PII / Anonymize", self._run_anonymize,
                      width=28).pack(padx=14, pady=(0, 14))

        self._build_preview_area(right, "anon")

    def _build_restore_tab(self):
        tab = self._tab_restore
        left = tk.Frame(tab, bg=DARK_BG, width=340)
        left.pack(side="left", fill="y", padx=(12, 6), pady=12)
        left.pack_propagate(False)
        right = tk.Frame(tab, bg=DARK_BG)
        right.pack(side="left", fill="both", expand=True, padx=(0, 12), pady=12)

        panel = tk.Frame(left, bg=PANEL_BG, bd=0)
        panel.pack(fill="both", expand=True, pady=(0, 8))

        section_label(panel, "RESTORE MODE").pack(anchor="w", padx=14, pady=(14, 4))

        self._restore_mode = tk.StringVar(value="original")
        mode_frame = tk.Frame(panel, bg=PANEL_BG)
        mode_frame.pack(fill="x", padx=14, pady=(0, 8))
        tk.Radiobutton(mode_frame, text="Restore original PII",
                       variable=self._restore_mode, value="original",
                       bg=PANEL_BG, fg=TEXT_MAIN, selectcolor=ENTRY_BG,
                       font=("Segoe UI", 9), activebackground=PANEL_BG,
                       command=self._on_restore_mode_change).pack(anchor="w")
        tk.Radiobutton(mode_frame, text="Restore translated PII",
                       variable=self._restore_mode, value="translated",
                       bg=PANEL_BG, fg=TEXT_MAIN, selectcolor=ENTRY_BG,
                       font=("Segoe UI", 9), activebackground=PANEL_BG,
                       command=self._on_restore_mode_change).pack(anchor="w")

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=(4, 10))
        section_label(panel, "INPUT / OUTPUT").pack(anchor="w", padx=14, pady=(0, 4))

        self._rest_input  = tk.StringVar()
        self._rest_map    = tk.StringVar()
        self._rest_transmap = tk.StringVar()
        self._rest_output = tk.StringVar()

        entry_row(panel, "Anonymized document", self._rest_input,
                  self._browse_rest_input).pack(fill="x", padx=14, pady=3)
        entry_row(panel, "Mapping file", self._rest_map,
                  self._browse_rest_map).pack(fill="x", padx=14, pady=3)

        # Translation map row — shown only in translated mode (not packed initially)
        self._rest_transmap_row = entry_row(panel, "Translation map", self._rest_transmap,
                                            self._browse_rest_transmap)
        # Output row — stored as reference for before= positioning
        self._rest_output_row = entry_row(panel, "Restored output", self._rest_output,
                                          self._browse_rest_out)
        self._rest_output_row.pack(fill="x", padx=14, pady=3)

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)

        info_box = tk.Frame(panel, bg=ENTRY_BG, bd=0)
        info_box.pack(fill="x", padx=14, pady=(0, 10))
        self._restore_info_lbl = tk.Label(info_box,
                 text="Restore original PII:\n\n"
                      "1. Anonymize your document\n"
                      "2. Send to cloud AI / translator\n"
                      "3. Save the AI response to a file\n"
                      "4. Load it as 'Anonymized document'\n"
                      "5. Load the original mapping.json\n"
                      "6. Click Restore\n\n"
                      "For translated documents, switch\n"
                      "to 'Restore translated PII' mode.",
                 bg=ENTRY_BG, fg=TEXT_DIM, font=("Segoe UI", 9),
                 justify="left", padx=12, pady=10)
        self._restore_info_lbl.pack()

        self._restore_btn = styled_button(panel, "\U0001f513  Restore Original PII",
                                          self._run_restore, width=28)
        self._restore_btn.pack(padx=14, pady=(0, 14))

        self._build_preview_area(right, "rest")

    def _on_restore_mode_change(self):
        """Show/hide the translation map row and update button label."""
        mode = self._restore_mode.get()
        if mode == "translated":
            # Insert the translation map row after the mapping file row
            self._rest_transmap_row.pack(fill="x", padx=14, pady=3,
                                         before=self._rest_output_row)
            self._restore_btn.config(text="\U0001f310  Restore Translated PII")
            self._restore_info_lbl.config(
                text="Restore translated PII:\n\n"
                     "1. Anonymize your document\n"
                     "2. Review Translation Map tab\n"
                     "3. Translate the anonymized doc\n"
                     "   (DeepL, ChatGPT, etc.)\n"
                     "4. Load translated doc here\n"
                     "5. Load the translation map\n"
                     "6. Click Restore\n\n"
                     "Result: translated document\n"
                     "with correct PII in target language.")
        else:
            self._rest_transmap_row.pack_forget()
            self._restore_btn.config(text="\U0001f513  Restore Original PII")
            self._restore_info_lbl.config(
                text="Restore original PII:\n\n"
                     "1. Anonymize your document\n"
                     "2. Send to cloud AI / translator\n"
                     "3. Save the AI response to a file\n"
                     "4. Load it as 'Anonymized document'\n"
                     "5. Load the original mapping.json\n"
                     "6. Click Restore\n\n"
                     "For translated documents, switch\n"
                     "to 'Restore translated PII' mode.")

    def _browse_rest_transmap(self):
        path = filedialog.askopenfilename(
            title="Select translation map file",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")])
        if path:
            self._rest_transmap.set(path)

    # -----------------------------------------------------------------------
    #  Translation Map tab
    # -----------------------------------------------------------------------

    def _build_translation_map_tab(self):
        tab = self._tab_transmap
        left = tk.Frame(tab, bg=DARK_BG, width=340)
        left.pack(side="left", fill="y", padx=(12, 6), pady=12)
        left.pack_propagate(False)
        right = tk.Frame(tab, bg=DARK_BG)
        right.pack(side="left", fill="both", expand=True, padx=(0, 12), pady=12)

        panel = tk.Frame(left, bg=PANEL_BG)
        panel.pack(fill="both", expand=True)

        section_label(panel, "TRANSLATION MAP FILE").pack(anchor="w", padx=14, pady=(14, 4))

        self._transmap_path = tk.StringVar()
        entry_row(panel, "Map file", self._transmap_path,
                  self._browse_transmap).pack(fill="x", padx=14, pady=3)

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=12)

        info_box = tk.Frame(panel, bg=ENTRY_BG)
        info_box.pack(fill="x", padx=14, pady=(0, 10))
        tk.Label(info_box,
                 text="How to use:\n\n"
                      "1. Anonymize a document — the\n"
                      "   translation map is auto-loaded.\n\n"
                      "2. Review the table on the right.\n"
                      "   Edit any \"Translated\" cell that\n"
                      "   needs correction.\n\n"
                      "3. Click Save Map.\n\n"
                      "4. Translate the anonymized doc\n"
                      "   (DeepL, ChatGPT, etc.).\n\n"
                      "5. In the Restore tab, load the\n"
                      "   translated doc + this map file.",
                 bg=ENTRY_BG, fg=TEXT_DIM, font=("Segoe UI", 9),
                 justify="left", padx=12, pady=10).pack()

        styled_button(panel, "\U0001f4be  Save Map", self._save_transmap,
                      width=28).pack(padx=14, pady=(0, 4))
        styled_button(panel, "\U0001f4c2  Load Map File", self._browse_transmap,
                      bg=PANEL_BG, width=28).pack(padx=14, pady=(0, 14))

        # Right side: editable table
        hdr = tk.Frame(right, bg=DARK_BG)
        hdr.pack(fill="x", pady=(0, 6))
        section_label(hdr, "PII TRANSLATION TABLE").pack(side="left", padx=4)
        tk.Label(hdr, text="Double-click a \"Translated\" cell to edit",
                 bg=DARK_BG, fg=TEXT_DIM, font=("Segoe UI", 8)).pack(side="left", padx=12)

        cols = ("Placeholder", "Entity", "Original", "Translated")
        self._transmap_tv = ttk.Treeview(right, columns=cols, show="headings", height=22)
        style = ttk.Style()
        style.configure("Treeview", background=ENTRY_BG, foreground=TEXT_MAIN,
                        fieldbackground=ENTRY_BG, rowheight=26, font=("Consolas", 9))
        col_widths = {"Placeholder": 160, "Entity": 120, "Original": 220, "Translated": 220}
        for col in cols:
            self._transmap_tv.heading(col, text=col)
            self._transmap_tv.column(col, width=col_widths[col], anchor="w")
        vsb = ttk.Scrollbar(right, orient="vertical", command=self._transmap_tv.yview)
        self._transmap_tv.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._transmap_tv.pack(fill="both", expand=True)
        self._transmap_tv.bind("<Double-1>", self._transmap_edit_cell)

        # Tag for rows with empty translation (needs attention)
        self._transmap_tv.tag_configure("needs_review", foreground=WARNING)
        self._transmap_tv.tag_configure("ok", foreground=SUCCESS)

        btn_row = tk.Frame(right, bg=DARK_BG)
        btn_row.pack(fill="x", pady=(6, 0))
        tk.Button(btn_row, text="Auto-fill Suggestions",
                  command=self._transmap_autofill,
                  bg=ACCENT, fg="white", font=("Segoe UI", 9, "bold"),
                  relief="flat", cursor="hand2", padx=10, pady=6).pack(side="left")
        tk.Button(btn_row, text="Clear Translated Column",
                  command=self._transmap_clear_translated,
                  bg=BORDER, fg=TEXT_MAIN, font=("Segoe UI", 9),
                  relief="flat", cursor="hand2", padx=10, pady=6).pack(side="left", padx=(8, 0))

        # Internal state
        self._transmap_data: Dict[str, dict] = {}  # ph -> {original, translated, entity}
        self._transmap_file_path: Optional[str] = None

    def _browse_transmap(self):
        path = filedialog.askopenfilename(
            title="Select translation map file",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")])
        if path:
            self._load_translation_map_file(path)

    def _load_translation_map_file(self, path: str):
        """Load a *_mapping_translated.json into the editor table."""
        try:
            with open(path, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            # Accept both formats:
            #   new: {ph: {original, translated, entity}}
            #   legacy plain mapping: {ph: original_str}
            normalized = {}
            for ph, val in data.items():
                if isinstance(val, dict):
                    normalized[ph] = {
                        "original":   val.get("original", ""),
                        "translated": val.get("translated", ""),
                        "entity":     val.get("entity", "CUSTOM"),
                    }
                else:
                    # Plain mapping — generate translation suggestions
                    plain_map = {ph2: v2 for ph2, v2 in data.items() if isinstance(v2, str)}
                    normalized = generate_translation_map(plain_map)
                    break
            self._transmap_data = normalized
            self._transmap_file_path = path
            self._transmap_path.set(path)
            self._refresh_transmap_table()
            logger.info("Translation map loaded: %s (%d entries)", path, len(normalized))
        except Exception as exc:
            logger.error("Failed to load translation map: %s", exc, exc_info=True)
            messagebox.showerror("Load Error", f"Could not load translation map:\n{exc}")

    def _refresh_transmap_table(self):
        """Repopulate the Treeview from self._transmap_data."""
        for row in self._transmap_tv.get_children():
            self._transmap_tv.delete(row)
        for ph, info in self._transmap_data.items():
            orig = info.get("original", "")
            trans = info.get("translated", "")
            entity = info.get("entity", "")
            tag = "ok" if trans and trans != "?" else "needs_review"
            self._transmap_tv.insert("", "end",
                values=(ph, entity, orig, trans), tags=(tag,))

    def _transmap_edit_cell(self, event):
        """Allow inline editing of the Translated column on double-click."""
        region = self._transmap_tv.identify_region(event.x, event.y)
        if region != "cell":
            return
        col_id = self._transmap_tv.identify_column(event.x)
        col_idx = int(col_id.replace("#", "")) - 1  # 0-based
        if col_idx != 3:  # Only allow editing column 4 (Translated)
            return
        item = self._transmap_tv.identify_row(event.y)
        if not item:
            return
        # Get cell bounding box
        x, y, w, h = self._transmap_tv.bbox(item, col_id)
        current_val = self._transmap_tv.item(item)["values"][3]
        # Create a temporary Entry widget over the cell
        edit_var = tk.StringVar(value=str(current_val))
        entry = tk.Entry(self._transmap_tv, textvariable=edit_var,
                         bg=ENTRY_BG, fg=TEXT_MAIN, insertbackground=TEXT_MAIN,
                         relief="flat", font=("Consolas", 9), bd=2)
        entry.place(x=x, y=y, width=w, height=h)
        entry.focus_set()
        entry.select_range(0, "end")

        def _commit(event=None):
            new_val = edit_var.get().strip()
            vals = list(self._transmap_tv.item(item)["values"])
            vals[3] = new_val
            self._transmap_tv.item(item, values=vals)
            # Update internal data
            ph = vals[0]
            if ph in self._transmap_data:
                self._transmap_data[ph]["translated"] = new_val
            tag = "ok" if new_val and new_val != "?" else "needs_review"
            self._transmap_tv.item(item, tags=(tag,))
            entry.destroy()

        entry.bind("<Return>", _commit)
        entry.bind("<FocusOut>", _commit)
        entry.bind("<Escape>", lambda e: entry.destroy())

    def _save_transmap(self):
        """Save the current translation map data back to the JSON file."""
        if not self._transmap_data:
            messagebox.showwarning("Nothing to Save", "No translation map is loaded.")
            return
        path = self._transmap_file_path
        if not path:
            path = filedialog.asksaveasfilename(
                title="Save translation map",
                defaultextension=".json",
                filetypes=[("JSON files", "*.json")])
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as fh:
                json.dump(self._transmap_data, fh, indent=4, ensure_ascii=False)
            self._transmap_file_path = path
            self._transmap_path.set(path)
            logger.info("Translation map saved: %s", path)
            messagebox.showinfo("Saved",
                f"Translation map saved:\n{path}\n\n"
                f"Use this file in the Restore tab to de-anonymize a translated document.")
        except Exception as exc:
            logger.error("Failed to save translation map: %s", exc, exc_info=True)
            messagebox.showerror("Save Error", f"Could not save:\n{exc}")

    def _transmap_autofill(self):
        """Re-run transliteration suggestions for all empty Translated cells."""
        if not self._transmap_data:
            messagebox.showwarning("Nothing Loaded", "Load a translation map first.")
            return
        # Rebuild suggestions only for empty/? entries
        plain = {ph: info["original"] for ph, info in self._transmap_data.items()}
        suggestions = generate_translation_map(plain)
        filled = 0
        for ph, info in self._transmap_data.items():
            if not info.get("translated") or info["translated"] == "?":
                suggestion = suggestions.get(ph, {}).get("translated", "")
                if suggestion and suggestion != "?":
                    self._transmap_data[ph]["translated"] = suggestion
                    filled += 1
        self._refresh_transmap_table()
        messagebox.showinfo("Auto-fill Complete",
            f"Filled {filled} empty translation(s) with suggestions.\n"
            "Please review and correct any that are wrong.")

    def _transmap_clear_translated(self):
        """Clear all translated values (reset to empty for manual entry)."""
        if not self._transmap_data:
            return
        for ph in self._transmap_data:
            self._transmap_data[ph]["translated"] = ""
        self._refresh_transmap_table()

    def _build_about_tab(self):
        tab = self._tab_about
        frame = tk.Frame(tab, bg=DARK_BG)
        frame.pack(expand=True)

        tk.Label(frame, text=f"Legal Document PII Anonymizer & Restorer  v{APP_VERSION}",
                 bg=DARK_BG, fg=TEXT_MAIN, font=("Segoe UI", 16, "bold")).pack(pady=(40, 0))
        tk.Label(frame, text="Protect attorney-client privilege before using cloud AI",
                 bg=DARK_BG, fg=TEXT_DIM, font=("Segoe UI", 10)).pack(pady=(4, 8))

        # Check for updates button
        update_row = tk.Frame(frame, bg=DARK_BG)
        update_row.pack(pady=(0, 16))
        self._update_status_lbl = tk.Label(
            update_row, text=f"Version {APP_VERSION}  \u2014  click to check for updates",
            bg=DARK_BG, fg=TEXT_DIM, font=("Segoe UI", 9)
        )
        self._update_status_lbl.pack(side="left", padx=(0, 12))
        tk.Button(
            update_row, text="\u21ba  Check for Updates",
            bg=PANEL_BG, fg=TEXT_MAIN, font=("Segoe UI", 9),
            relief="flat", cursor="hand2", padx=10,
            command=self._manual_check_update,
        ).pack(side="left")

        info = [
            ("Engine",      "Microsoft Presidio + spaCy (en_core_web_sm + xx_ent_wiki_sm)"),
            ("Languages",   "English + Hebrew — auto-detected per document"),
            ("Hebrew NER",  "700+ Israeli first names, 200+ surnames, 150+ locations (dictionary)"),
            ("Israeli PII", "Teudat Zehut (Luhn-10), phone (+972/05x/07x), IBAN IL"),
            ("Entities",    f"{len(set(ENTITY_LABELS.values()))} PII types detected"),
            ("Formats",     ".txt  |  .docx  |  .pdf  (input)    .txt  |  .docx  (output)"),
            ("Encoding",    "UTF-8, Windows-1255, ISO-8859-8, CP1255"),
            ("Privacy",     "100% local — no data leaves your machine"),
        ]
        for key, val in info:
            row = tk.Frame(frame, bg=PANEL_BG, pady=8, padx=20)
            row.pack(fill="x", padx=60, pady=3)
            tk.Label(row, text=f"{key}:", bg=PANEL_BG, fg=ACCENT,
                     font=("Segoe UI", 9, "bold"), width=14, anchor="w").pack(side="left")
            tk.Label(row, text=val, bg=PANEL_BG, fg=TEXT_MAIN, font=("Segoe UI", 9)).pack(side="left")

    def _build_preview_area(self, parent, prefix):
        top = tk.Frame(parent, bg=DARK_BG)
        top.pack(fill="both", expand=True)

        left_frame  = tk.Frame(top, bg=DARK_BG)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 4))
        right_frame = tk.Frame(top, bg=DARK_BG)
        right_frame.pack(side="left", fill="both", expand=True, padx=(4, 0))

        lbl_in  = "Original Document"  if prefix == "anon" else "Anonymized Document"
        lbl_out = "Anonymized Output"  if prefix == "anon" else "Restored Document"

        for frame, lbl, attr in [
            (left_frame,  lbl_in,  f"_{prefix}_txt_in"),
            (right_frame, lbl_out, f"_{prefix}_txt_out"),
        ]:
            header = tk.Frame(frame, bg=PANEL_BG)
            header.pack(fill="x")
            tk.Label(header, text=lbl, bg=PANEL_BG, fg=ACCENT,
                     font=("Segoe UI", 9, "bold"), pady=6, padx=10).pack(side="left")
            if attr.endswith("_out"):
                tk.Button(header, text="Copy", bg=BORDER, fg=TEXT_MAIN,
                          font=("Segoe UI", 8), relief="flat", cursor="hand2",
                          padx=6, command=lambda a=attr: self._copy_text(a)
                          ).pack(side="right", padx=6, pady=4)
                tk.Button(header, text="Save As...", bg=BORDER, fg=TEXT_MAIN,
                          font=("Segoe UI", 8), relief="flat", cursor="hand2",
                          padx=6, command=lambda a=attr: self._save_text(a)
                          ).pack(side="right", padx=(0, 4), pady=4)

            txt = scrolledtext.ScrolledText(
                frame, bg=ENTRY_BG, fg=TEXT_MAIN, insertbackground=TEXT_MAIN,
                font=("Consolas", 9), relief="flat", wrap="word",
                selectbackground=ACCENT, selectforeground="white",
            )
            txt.pack(fill="both", expand=True)
            setattr(self, attr, txt)

        out_widget = getattr(self, f"_{prefix}_txt_out")
        for label, colour in ENTITY_COLOURS.items():
            out_widget.tag_configure(label, foreground=colour, font=("Consolas", 9, "bold"))
        out_widget.tag_configure("OTHER", foreground=TAG_OTHER, font=("Consolas", 9, "bold"))

        tbl_frame = tk.Frame(parent, bg=PANEL_BG, height=160)
        tbl_frame.pack(fill="x", pady=(8, 0))
        tbl_frame.pack_propagate(False)

        tk.Label(tbl_frame, text="Detection Log", bg=PANEL_BG, fg=ACCENT,
                 font=("Segoe UI", 9, "bold"), pady=6, padx=10).pack(anchor="w")

        cols = ("Placeholder", "Original Value", "Entity Type", "Confidence", "Lang")
        tv = ttk.Treeview(tbl_frame, columns=cols, show="headings", height=4)
        style = ttk.Style()
        style.configure("Treeview", background=ENTRY_BG, foreground=TEXT_MAIN,
                         fieldbackground=ENTRY_BG, rowheight=22, font=("Consolas", 8))
        style.configure("Treeview.Heading", background=PANEL_BG, foreground=ACCENT,
                         font=("Segoe UI", 8, "bold"), relief="flat")
        style.map("Treeview", background=[("selected", ACCENT)])

        col_widths = {"Placeholder": 130, "Original Value": 200,
                      "Entity Type": 120, "Confidence": 80, "Lang": 50}
        for col in cols:
            tv.heading(col, text=col)
            tv.column(col, width=col_widths.get(col, 120), anchor="w")

        vsb = ttk.Scrollbar(tbl_frame, orient="vertical", command=tv.yview)
        tv.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        tv.pack(fill="both", expand=True, padx=6, pady=(0, 6))
        setattr(self, f"_{prefix}_table", tv)

    def _browse_input(self):
        path = filedialog.askopenfilename(
            title="Select input document",
            filetypes=[("Documents", "*.txt *.docx *.pdf"), ("All files", "*.*")])
        if path:
            self._anon_input.set(path)
            base = os.path.splitext(path)[0]
            self._anon_output.set(base + "_anonymized.txt")
            self._anon_map.set(base + "_mapping.json")
            self._load_preview(path, self._anon_txt_in)
            # Auto-load project custom PII from the same folder
            folder = os.path.dirname(path)
            self._load_project_pii_into_editor(folder)

    def _browse_anon_out(self):
        path = filedialog.asksaveasfilename(title="Save anonymized document",
            defaultextension=".txt", filetypes=[("Text file", "*.txt"), ("Word document", "*.docx")])
        if path: self._anon_output.set(path)

    def _browse_anon_map(self):
        path = filedialog.asksaveasfilename(title="Save mapping file",
            defaultextension=".json", filetypes=[("JSON file", "*.json")])
        if path: self._anon_map.set(path)

    def _browse_rest_input(self):
        path = filedialog.askopenfilename(title="Select anonymized document",
            filetypes=[("Documents", "*.txt *.docx"), ("All files", "*.*")])
        if path:
            self._rest_input.set(path)
            base = os.path.splitext(path)[0]
            self._rest_output.set(base + "_restored.txt")
            self._load_preview(path, self._rest_txt_in)

    def _browse_rest_map(self):
        path = filedialog.askopenfilename(title="Select mapping file",
            filetypes=[("JSON file", "*.json"), ("All files", "*.*")])
        if path: self._rest_map.set(path)

    def _browse_rest_out(self):
        path = filedialog.asksaveasfilename(title="Save restored document",
            defaultextension=".txt", filetypes=[("Text file", "*.txt"), ("Word document", "*.docx")])
        if path: self._rest_output.set(path)

    def _load_preview(self, path: str, widget: scrolledtext.ScrolledText):
        try:
            text = read_file(path)
            widget.config(state="normal")
            widget.delete("1.0", "end")
            widget.insert("1.0", text)
            if is_rtl(text):
                try: widget.config(justify="right")
                except Exception: pass
            widget.config(state="disabled")
        except Exception as exc:
            messagebox.showerror("Read Error", str(exc))

    def _set_output_text(self, widget: scrolledtext.ScrolledText, text: str, detections=None):
        widget.config(state="normal")
        widget.delete("1.0", "end")
        widget.insert("1.0", text)
        if is_rtl(text):
            try: widget.config(justify="right")
            except Exception: pass
        if detections:
            self._highlight_placeholders(widget, text, detections)
        widget.config(state="disabled")

    def _highlight_placeholders(self, widget, text: str, detections):
        for ph_match in re.finditer(r"\{\{([A-Z_]+)_\d+\}\}", text):
            label = ph_match.group(1)
            tag   = label if label in ENTITY_COLOURS else "OTHER"
            widget.tag_add(tag,
                           f"1.0 + {ph_match.start()} chars",
                           f"1.0 + {ph_match.end()} chars")

    def _copy_text(self, attr: str):
        text = getattr(self, attr).get("1.0", "end-1c")
        self.clipboard_clear()
        self.clipboard_append(text)
        self._set_status("Copied to clipboard.", SUCCESS)

    def _save_text(self, attr: str):
        text = getattr(self, attr).get("1.0", "end-1c")
        path = filedialog.asksaveasfilename(defaultextension=".txt",
            filetypes=[("Text file", "*.txt"), ("Word document", "*.docx")])
        if path:
            try:
                write_file(path, text)
                self._set_status(f"Saved to {path}", SUCCESS)
            except Exception as exc:
                messagebox.showerror("Save Error", str(exc))

    def _select_all_entities(self):
        for v in self._entity_vars.values(): v.set(True)

    def _clear_all_entities(self):
        for v in self._entity_vars.values(): v.set(False)

    def _set_status(self, msg: str, colour: str = TEXT_MAIN):
        self._status_lbl.config(text=msg, fg=colour)
        self._bar_lbl.config(text=msg)

    def _start_spinner(self, msg: str):
        self._bar_lbl.config(text=msg)
        self._progress.start(12)

    def _stop_spinner(self):
        self._progress.stop()

    def _populate_table(self, tv: ttk.Treeview, detections):
        for row in tv.get_children(): tv.delete(row)
        seen = set()
        for d in sorted(detections, key=lambda x: x["start"]):
            key = (d["placeholder"], d["original"])
            if key not in seen:
                seen.add(key)
                tv.insert("", "end", values=(
                    d["placeholder"], d["original"], d["label"],
                    f"{d['score']:.2f}", d.get("lang", "en").upper(),
                ))

    def _run_anonymize(self):
        if not self._engine_ready:
            messagebox.showwarning("Engine Loading", "Please wait for the NLP engine to finish loading.")
            return
        inp  = self._anon_input.get().strip()
        out  = self._anon_output.get().strip()
        mapf = self._anon_map.get().strip()
        if not inp:
            messagebox.showwarning("Missing Input", "Please select an input document.")
            return
        if not out:
            messagebox.showwarning("Missing Output", "Please specify an output file path.")
            return
        if not mapf:
            messagebox.showwarning("Missing Mapping", "Please specify a mapping file path.")
            return

        selected_entities = [e for e, v in self._entity_vars.items() if v.get()]
        confidence = self._confidence.get()
        self._start_spinner("Anonymizing document...")
        self._set_status("Anonymizing...", WARNING)

        # Collect custom PII entries from the project editor
        custom_entries = self._get_custom_pii_entries()

        def task():
            """Phase 1: Read file and run NLP detection in background thread."""
            try:
                logger.info("Anonymize task started: input=%s", os.path.basename(inp))
                text = read_file(inp)
                logger.debug("File read: %d chars from %s", len(text), inp)
                engine = PIIEngine.get()
                anon_text, mapping, detections = engine.anonymize(
                    text, confidence=confidence, entities=selected_entities)

                # --- Hebrew ambiguity review ---
                # If the document is Hebrew, find ambiguous candidates not caught by NLP
                lang = detect_language(text)
                if lang == "he":
                    already_mapped = set(mapping.values())
                    candidates = find_hebrew_ambiguous_candidates(text, already_mapped)
                    logger.debug("Hebrew ambiguity candidates: %d", len(candidates))
                else:
                    candidates = []

                # Post to UI thread for optional review dialog, then finalize
                self.after(0, lambda: self._on_nlp_done(
                    text, anon_text, mapping, detections,
                    candidates, custom_entries, out, mapf
                ))
            except Exception as exc:
                logger.error("Anonymize task failed: %s", exc, exc_info=True)
                self.after(0, lambda: self._on_error(str(exc)))

        threading.Thread(target=task, daemon=True).start()

    def _on_nlp_done(self, original_text, anon_text, mapping, detections,
                     candidates, custom_entries, out, mapf):
        """
        Phase 2 (UI thread): optionally show Hebrew review dialog,
        apply approved candidates + custom entries, then write output files.
        """
        # Show Hebrew review dialog if there are ambiguous candidates
        approved_hebrew = []
        if candidates:
            self._stop_spinner()  # pause spinner while user reviews
            self._set_status(
                f"Hebrew review: {len(candidates)} ambiguous term(s) found — please confirm.",
                WARNING
            )
            dlg = HebrewReviewDialog(self, candidates)
            self.wait_window(dlg)  # blocks UI until dialog is closed
            approved_hebrew = dlg.approved
            self._start_spinner("Finalizing anonymization...")
            self._set_status("Finalizing...", WARNING)

        try:
            # Rebuild helper dicts from existing mapping
            entity_counts: Dict[str, int] = {}
            value_to_ph: Dict[str, str] = {v: k for k, v in mapping.items()}
            for ph in list(mapping.keys()):
                m = re.match(r'\{\{([A-Z_]+)_(\d+)\}\}', ph)
                if m:
                    lbl, num = m.group(1), int(m.group(2))
                    entity_counts[lbl] = max(entity_counts.get(lbl, 0), num)

            # Apply Hebrew-approved candidates
            if approved_hebrew:
                logger.info("Applying %d Hebrew-approved candidates", len(approved_hebrew))
                anon_text = apply_custom_pii(
                    anon_text, approved_hebrew,
                    mapping, entity_counts, value_to_ph, detections)

            # Apply manual custom PII entries
            if custom_entries:
                logger.debug("Applying %d custom PII entries", len(custom_entries))
                anon_text = apply_custom_pii(
                    anon_text, custom_entries,
                    mapping, entity_counts, value_to_ph, detections)

            write_file(out, anon_text)
            logger.debug("Anonymized output written: %s", out)
            with open(mapf, "w", encoding="utf-8") as fh:
                json.dump(mapping, fh, indent=4, ensure_ascii=False)
            logger.info("Mapping file written: %s (%d entries)", mapf, len(mapping))
            # Generate and write the translation map alongside the original mapping
            trans_mapf = _get_translation_map_path(mapf)
            try:
                trans_map = generate_translation_map(mapping)
                with open(trans_mapf, "w", encoding="utf-8") as fh:
                    json.dump(trans_map, fh, indent=4, ensure_ascii=False)
                logger.info("Translation map written: %s", trans_mapf)
            except Exception as exc:
                logger.warning("Could not write translation map: %s", exc)
                trans_mapf = None
            self._on_anonymize_done(anon_text, mapping, detections, out, mapf, trans_mapf)
        except Exception as exc:
            logger.error("Finalize anonymize failed: %s", exc, exc_info=True)
            self._on_error(str(exc))

    def _on_anonymize_done(self, anon_text, mapping, detections, out, mapf, trans_mapf=None):
        self._stop_spinner()
        self._set_output_text(self._anon_txt_out, anon_text, detections)
        self._populate_table(self._anon_table, detections)
        n_unique = len(mapping)
        n_total  = len(detections)
        self._set_status(f"Done — {n_unique} unique PII items replaced ({n_total} total).", SUCCESS)
        # Auto-load the translation map into the Translation Map tab
        if trans_mapf and os.path.exists(trans_mapf):
            self.after(0, lambda: self._load_translation_map_file(trans_mapf))
        trans_note = (
            f"\nTranslation map: {os.path.basename(trans_mapf) if trans_mapf else 'not generated'}\n"
            f"  → Review in the \"Translation Map\" tab before restoring a translated document."
        ) if trans_mapf else ""
        messagebox.showinfo("Anonymization Complete",
            f"Anonymization complete!\n\n"
            f"  Unique PII items replaced : {n_unique}\n"
            f"  Total occurrences         : {n_total}\n\n"
            f"Anonymized document: {out}\n"
            f"Mapping file: {mapf}{trans_note}\n\n"
            f"Safe to send to cloud AI or translator.")

    def _run_restore(self):
        inp  = self._rest_input.get().strip()
        mapf = self._rest_map.get().strip()
        out  = self._rest_output.get().strip()
        mode = self._restore_mode.get()
        trans_mapf = self._rest_transmap.get().strip() if mode == "translated" else None

        if not inp:
            messagebox.showwarning("Missing Input", "Please select an anonymized document.")
            return
        if not mapf:
            messagebox.showwarning("Missing Mapping", "Please select a mapping file.")
            return
        if mode == "translated" and not trans_mapf:
            messagebox.showwarning("Missing Translation Map",
                "Please select a translation map file (.json) in the 'Translation map' field.\n\n"
                "This file is generated automatically when you anonymize a document.\n"
                "Review and edit it in the Translation Map tab first.")
            return
        if not out:
            messagebox.showwarning("Missing Output", "Please specify an output file path.")
            return

        self._start_spinner("Restoring PII...")
        self._set_status("Restoring...", WARNING)

        def task():
            try:
                logger.info("Restore task started: mode=%s input=%s", mode, os.path.basename(inp))
                text = read_file(inp)
                logger.debug("Anonymized file read: %d chars", len(text))

                if mode == "translated" and trans_mapf:
                    # Load the translation map and build a flat {placeholder: translated_value} dict
                    with open(trans_mapf, "r", encoding="utf-8") as fh:
                        trans_data = json.load(fh)
                    # Build flat restore mapping: ph -> translated value
                    restore_mapping = {}
                    for ph, info in trans_data.items():
                        if isinstance(info, dict):
                            translated = info.get("translated", "").strip()
                            original   = info.get("original", "")
                            # Use translated if available, fall back to original
                            restore_mapping[ph] = translated if translated and translated != "?" else original
                        else:
                            restore_mapping[ph] = str(info)
                    logger.debug("Translation mapping loaded: %d entries from %s",
                                 len(restore_mapping), trans_mapf)
                    # Also load the original mapping for any placeholders not in trans map
                    with open(mapf, "r", encoding="utf-8") as fh:
                        orig_mapping = json.load(fh)
                    for ph, orig in orig_mapping.items():
                        if ph not in restore_mapping:
                            restore_mapping[ph] = orig
                    display_mapping = restore_mapping
                else:
                    # Standard restore with original mapping
                    with open(mapf, "r", encoding="utf-8") as fh:
                        restore_mapping = json.load(fh)
                    display_mapping = restore_mapping
                    logger.debug("Mapping loaded: %d entries from %s", len(restore_mapping), mapf)

                restored = PIIEngine.restore(text, restore_mapping)
                write_file(out, restored)
                logger.info("Restore complete: output written to %s", out)
                self.after(0, lambda: self._on_restore_done(restored, display_mapping, out, mode))
            except Exception as exc:
                logger.error("Restore task failed: %s", exc, exc_info=True)
                self.after(0, lambda: self._on_error(str(exc)))

        threading.Thread(target=task, daemon=True).start()

    def _on_restore_done(self, restored, mapping, out, mode="original"):
        self._stop_spinner()
        self._set_output_text(self._rest_txt_out, restored)
        fake_detections = [
            {"placeholder": ph, "original": orig,
             "label": ph.strip("{}").rsplit("_", 1)[0],
             "score": 1.0, "start": 0, "end": 0, "lang": ""}
            for ph, orig in mapping.items()
        ]
        self._populate_table(self._rest_table, fake_detections)
        self._set_status(f"Done — {len(mapping)} placeholders restored.", SUCCESS)
        mode_label = "translated" if mode == "translated" else "original"
        messagebox.showinfo("Restoration Complete",
            f"Restoration complete!\n\n"
            f"  Mode                  : {mode_label} PII\n"
            f"  Placeholders restored : {len(mapping)}\n\n"
            f"Restored document saved to:\n  {out}")

    def _manual_check_update(self):
        """Called when the user clicks 'Check for Updates' in the About tab."""
        self._update_status_lbl.config(text="Checking for updates...", fg=WARNING)

        def _check():
            try:
                req = urllib.request.Request(
                    UPDATE_MANIFEST_URL,
                    headers={"User-Agent": f"PII-Processor/{APP_VERSION}"},
                )
                ctx = _make_ssl_context()
                with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
                    data = json.loads(resp.read().decode())
                latest  = data.get("version", "")
                url     = data.get("download_url", "")
                notes   = data.get("release_notes", "")
                logger.debug("Manual update check: latest=%s current=%s", latest, APP_VERSION)
                if latest and url and Updater._is_newer(latest):
                    self.after(0, lambda: self._update_status_lbl.config(
                        text=f"\u2b06 Update available: v{latest}  ({notes})",
                        fg=SUCCESS,
                    ))
                    self.after(0, lambda: self._updater._on_update_available(
                        latest, url, notes
                    ))
                elif latest:
                    self.after(0, lambda: self._update_status_lbl.config(
                        text=f"\u2714 You are on the latest version (v{APP_VERSION})",
                        fg=SUCCESS,
                    ))
                else:
                    self.after(0, lambda: self._update_status_lbl.config(
                        text="Could not read version manifest", fg=DANGER
                    ))
            except urllib.error.URLError as exc:
                reason = exc.reason if exc.reason else str(exc)
                err_msg = f"Network error: {reason}"
                logger.warning("Manual update check URLError: %s", reason)
                self.after(0, lambda m=err_msg: self._update_status_lbl.config(
                    text=m, fg=DANGER
                ))
            except Exception as exc:
                err_msg = f"Update check failed: {type(exc).__name__}: {exc}"
                logger.warning("Manual update check exception: %s", err_msg)
                self.after(0, lambda m=err_msg: self._update_status_lbl.config(
                    text=m, fg=DANGER
                ))

        threading.Thread(target=_check, daemon=True).start()

    # -----------------------------------------------------------------------
    #  Debug Log tab
    # -----------------------------------------------------------------------

    def _build_debug_tab(self):
        tab = self._tab_debug

        # ---- toolbar ----
        toolbar = tk.Frame(tab, bg=PANEL_BG)
        toolbar.pack(fill="x", padx=12, pady=(10, 0))

        section_label(toolbar, "DEBUG LOG").pack(side="left", padx=(4, 16))

        # Debug mode toggle
        self._debug_mode = tk.BooleanVar(value=False)
        def _toggle_debug():
            if self._debug_mode.get():
                MEM_LOG_HANDLER.setLevel(logging.DEBUG)
                _file_handler.setLevel(logging.DEBUG)
                logger.info("Debug mode ENABLED")
                self._debug_toggle_btn.config(text="U0001f41b Debug Mode: ON", fg=SUCCESS)
            else:
                MEM_LOG_HANDLER.setLevel(logging.INFO)
                _file_handler.setLevel(logging.INFO)
                logger.info("Debug mode DISABLED")
                self._debug_toggle_btn.config(text="U0001f41b Debug Mode: OFF", fg=TEXT_DIM)

        self._debug_toggle_btn = tk.Button(
            toolbar, text="U0001f41b Debug Mode: OFF",
            bg=PANEL_BG, fg=TEXT_DIM, relief="flat",
            font=("Segoe UI", 9), cursor="hand2",
            command=lambda: [self._debug_mode.set(not self._debug_mode.get()), _toggle_debug()]
        )
        self._debug_toggle_btn.pack(side="left", padx=4)

        tk.Label(toolbar, text="|", bg=PANEL_BG, fg=BORDER).pack(side="left", padx=4)

        tk.Button(toolbar, text="U0001f5d1  Clear",
                  bg=PANEL_BG, fg=TEXT_DIM, relief="flat",
                  font=("Segoe UI", 9), cursor="hand2",
                  command=self._debug_clear).pack(side="left", padx=4)

        tk.Button(toolbar, text="U0001f4cb  Copy Log",
                  bg=PANEL_BG, fg=TEXT_DIM, relief="flat",
                  font=("Segoe UI", 9), cursor="hand2",
                  command=self._debug_copy).pack(side="left", padx=4)

        tk.Button(toolbar, text="U0001f4c4  Open Log File",
                  bg=PANEL_BG, fg=TEXT_DIM, relief="flat",
                  font=("Segoe UI", 9), cursor="hand2",
                  command=self._debug_open_file).pack(side="left", padx=4)

        lp_lbl = tk.Label(toolbar, text=f"Log: {LOG_PATH}",
                          bg=PANEL_BG, fg=TEXT_DIM, font=("Segoe UI", 8))
        lp_lbl.pack(side="right", padx=8)

        # ---- log text widget ----
        self._debug_txt = scrolledtext.ScrolledText(
            tab, bg="#0d0d0d", fg="#cccccc",
            insertbackground="white",
            font=("Consolas", 9), relief="flat", wrap="word",
            selectbackground=ACCENT, selectforeground="white",
            state="disabled"
        )
        self._debug_txt.pack(fill="both", expand=True, padx=12, pady=(6, 12))

        # colour tags per level
        self._debug_txt.tag_configure("DEBUG",    foreground="#888888")
        self._debug_txt.tag_configure("INFO",     foreground="#cccccc")
        self._debug_txt.tag_configure("WARNING",  foreground=WARNING)
        self._debug_txt.tag_configure("ERROR",    foreground=DANGER)
        self._debug_txt.tag_configure("CRITICAL", foreground="#ff00ff")

        # Populate with records already in the buffer
        for rec in MEM_LOG_HANDLER.get_records():
            self._debug_append_record(rec)

        # Register live callback
        MEM_LOG_HANDLER.add_callback(lambda rec: self.after(0, lambda r=rec: self._debug_append_record(r)))

    def _debug_append_record(self, record: logging.LogRecord):
        """Append a single log record to the debug text widget."""
        try:
            msg = MEM_LOG_HANDLER.format(record)
            tag = record.levelname  # DEBUG / INFO / WARNING / ERROR / CRITICAL
            self._debug_txt.config(state="normal")
            self._debug_txt.insert("end", msg + "\n", tag)
            self._debug_txt.see("end")
            self._debug_txt.config(state="disabled")
        except Exception:
            pass

    def _debug_clear(self):
        self._debug_txt.config(state="normal")
        self._debug_txt.delete("1.0", "end")
        self._debug_txt.config(state="disabled")
        logger.info("Debug log display cleared by user")

    def _debug_copy(self):
        content = self._debug_txt.get("1.0", "end")
        self.clipboard_clear()
        self.clipboard_append(content)
        self._set_status("Debug log copied to clipboard.", SUCCESS)

    def _debug_open_file(self):
        """Open the log file in the system default text editor."""
        try:
            if sys.platform.startswith("win"):
                os.startfile(LOG_PATH)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", LOG_PATH])
            else:
                subprocess.Popen(["xdg-open", LOG_PATH])
        except Exception as exc:
            messagebox.showerror("Cannot Open Log",
                f"Could not open log file:\n{LOG_PATH}\n\nError: {exc}")

    def _on_error(self, msg: str):
        logger.error("UI error: %s", msg)
        self._stop_spinner()
        self._set_status(f"Error: {msg}", DANGER)
        messagebox.showerror("Error", msg)

    # -----------------------------------------------------------------------
    #  Custom PII tab
    # -----------------------------------------------------------------------

    def _build_custom_pii_tab(self):
        tab = self._tab_custom
        left = tk.Frame(tab, bg=DARK_BG, width=340)
        left.pack(side="left", fill="y", padx=(12, 6), pady=12)
        left.pack_propagate(False)
        right = tk.Frame(tab, bg=DARK_BG)
        right.pack(side="left", fill="both", expand=True, padx=(0, 12), pady=12)

        panel = tk.Frame(left, bg=PANEL_BG)
        panel.pack(fill="both", expand=True)

        section_label(panel, "PROJECT FOLDER").pack(anchor="w", padx=14, pady=(14, 4))

        self._custom_folder = tk.StringVar()
        folder_row = tk.Frame(panel, bg=PANEL_BG)
        folder_row.pack(fill="x", padx=14, pady=3)
        tk.Label(folder_row, text="Project folder", bg=PANEL_BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9), width=14, anchor="w").pack(side="left")
        tk.Entry(folder_row, textvariable=self._custom_folder, bg=ENTRY_BG, fg=TEXT_MAIN,
                 insertbackground=TEXT_MAIN, relief="flat", font=("Segoe UI", 9), bd=4
                 ).pack(side="left", fill="x", expand=True)
        tk.Button(folder_row, text="Browse", command=self._browse_custom_folder,
                  bg=BORDER, fg=TEXT_MAIN, font=("Segoe UI", 8), relief="flat",
                  cursor="hand2", padx=6).pack(side="left", padx=(4, 0))

        self._custom_pii_path_lbl = tk.Label(panel, text="", bg=PANEL_BG, fg=TEXT_DIM,
                                              font=("Segoe UI", 8), wraplength=300, justify="left")
        self._custom_pii_path_lbl.pack(anchor="w", padx=14, pady=(2, 0))

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)
        section_label(panel, "ADD NEW ENTRY").pack(anchor="w", padx=14, pady=(0, 6))

        self._new_pii_text  = tk.StringVar()
        self._new_pii_label = tk.StringVar(value="PERSON")

        entry_row(panel, "Text to hide", self._new_pii_text).pack(fill="x", padx=14, pady=3)

        lbl_row = tk.Frame(panel, bg=PANEL_BG)
        lbl_row.pack(fill="x", padx=14, pady=3)
        tk.Label(lbl_row, text="Entity label", bg=PANEL_BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9), width=14, anchor="w").pack(side="left")
        ttk.Combobox(lbl_row, textvariable=self._new_pii_label,
                     values=["PERSON", "LOCATION", "ORGANIZATION", "IL_ID",
                             "PHONE", "EMAIL", "DATE", "CUSTOM"],
                     state="normal", width=16, font=("Segoe UI", 9)).pack(side="left")

        styled_button(panel, "+ Add Entry", self._add_custom_pii_entry,
                      bg=SUCCESS, fg=DARK_BG, width=20).pack(padx=14, pady=(8, 4))

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=8)

        styled_button(panel, "\U0001f4be  Save to Project Folder", self._save_custom_pii,
                      width=28).pack(padx=14, pady=(0, 4))
        styled_button(panel, "\U0001f4c2  Load from Project Folder", self._load_custom_pii,
                      bg=PANEL_BG, width=28).pack(padx=14, pady=(0, 14))

        # Right side: list of current entries
        section_label(right, "CURRENT CUSTOM PII ENTRIES").pack(anchor="w", padx=4, pady=(4, 6))

        cols = ("Text to Hide", "Entity Label")
        self._custom_tv = ttk.Treeview(right, columns=cols, show="headings", height=20)
        style = ttk.Style()
        style.configure("Treeview", background=ENTRY_BG, foreground=TEXT_MAIN,
                         fieldbackground=ENTRY_BG, rowheight=24, font=("Consolas", 9))
        for col in cols:
            self._custom_tv.heading(col, text=col)
            self._custom_tv.column(col, width=280 if col == "Text to Hide" else 160, anchor="w")
        vsb = ttk.Scrollbar(right, orient="vertical", command=self._custom_tv.yview)
        self._custom_tv.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._custom_tv.pack(fill="both", expand=True)

        btn_row = tk.Frame(right, bg=DARK_BG)
        btn_row.pack(fill="x", pady=(6, 0))
        tk.Button(btn_row, text="\u274c  Remove Selected", command=self._remove_custom_pii_entry,
                  bg=DANGER, fg="white", font=("Segoe UI", 9, "bold"), relief="flat",
                  cursor="hand2", padx=10, pady=6).pack(side="left")
        tk.Button(btn_row, text="Clear All", command=self._clear_custom_pii,
                  bg=BORDER, fg=TEXT_MAIN, font=("Segoe UI", 9), relief="flat",
                  cursor="hand2", padx=10, pady=6).pack(side="left", padx=(8, 0))

    def _browse_custom_folder(self):
        folder = filedialog.askdirectory(title="Select project folder")
        if folder:
            self._custom_folder.set(folder)
            pii_path = get_project_pii_path(folder)
            self._custom_pii_path_lbl.config(
                text=f"PII file: {pii_path}",
                fg=SUCCESS if os.path.exists(pii_path) else TEXT_DIM)
            # Auto-load if file exists
            if os.path.exists(pii_path):
                self._load_project_pii_into_editor(folder)

    def _load_project_pii_into_editor(self, folder: str):
        """Load the project custom PII file into the editor table."""
        entries = load_project_pii(folder)
        if entries:
            self._custom_folder.set(folder)
            pii_path = get_project_pii_path(folder)
            self._custom_pii_path_lbl.config(
                text=f"PII file: {pii_path}", fg=SUCCESS)
            for row in self._custom_tv.get_children():
                self._custom_tv.delete(row)
            for e in entries:
                self._custom_tv.insert("", "end", values=(e.get("text", ""), e.get("label", "CUSTOM")))

    def _add_custom_pii_entry(self):
        text  = self._new_pii_text.get().strip()
        label = self._new_pii_label.get().strip().upper().replace(" ", "_") or "CUSTOM"
        if not text:
            messagebox.showwarning("Empty Entry", "Please enter the text you want to hide.")
            return
        # Check for duplicate
        for row in self._custom_tv.get_children():
            if self._custom_tv.item(row)["values"][0] == text:
                messagebox.showinfo("Duplicate", f"'{text}' is already in the list.")
                return
        self._custom_tv.insert("", "end", values=(text, label))
        self._new_pii_text.set("")

    def _remove_custom_pii_entry(self):
        selected = self._custom_tv.selection()
        if not selected:
            messagebox.showinfo("Nothing Selected", "Click a row to select it, then click Remove.")
            return
        for item in selected:
            self._custom_tv.delete(item)

    def _clear_custom_pii(self):
        if messagebox.askyesno("Clear All", "Remove all custom PII entries from the list?"):
            for row in self._custom_tv.get_children():
                self._custom_tv.delete(row)

    def _save_custom_pii(self):
        folder = self._custom_folder.get().strip()
        if not folder:
            folder = filedialog.askdirectory(title="Select project folder to save custom PII")
            if not folder:
                return
            self._custom_folder.set(folder)
        entries = self._get_custom_pii_entries()
        save_project_pii(folder, entries)
        pii_path = get_project_pii_path(folder)
        self._custom_pii_path_lbl.config(text=f"PII file: {pii_path}", fg=SUCCESS)
        self._set_status(f"Custom PII saved to {pii_path}", SUCCESS)
        messagebox.showinfo("Saved",
            f"Custom PII list saved:\n{pii_path}\n\n"
            f"{len(entries)} entries will be applied to all documents in this folder.")

    def _load_custom_pii(self):
        folder = self._custom_folder.get().strip()
        if not folder:
            folder = filedialog.askdirectory(title="Select project folder to load custom PII from")
            if not folder:
                return
            self._custom_folder.set(folder)
        pii_path = get_project_pii_path(folder)
        if not os.path.exists(pii_path):
            messagebox.showinfo("Not Found",
                f"No custom PII file found in:\n{folder}\n\n"
                f"(Expected: {CUSTOM_PII_FILENAME})")
            return
        self._load_project_pii_into_editor(folder)
        self._set_status(f"Loaded {len(load_project_pii(folder))} custom entries from {folder}", SUCCESS)

    def _get_custom_pii_entries(self) -> List[dict]:
        """Read current entries from the custom PII treeview."""
        entries = []
        for row in self._custom_tv.get_children():
            vals = self._custom_tv.item(row)["values"]
            if vals and str(vals[0]).strip():
                entries.append({"text": str(vals[0]), "label": str(vals[1])})
        return entries

    # -----------------------------------------------------------------------
    #  Batch processing tab
    # -----------------------------------------------------------------------

    def _build_batch_tab(self):
        tab = self._tab_batch
        left = tk.Frame(tab, bg=DARK_BG, width=340)
        left.pack(side="left", fill="y", padx=(12, 6), pady=12)
        left.pack_propagate(False)
        right = tk.Frame(tab, bg=DARK_BG)
        right.pack(side="left", fill="both", expand=True, padx=(0, 12), pady=12)

        panel = tk.Frame(left, bg=PANEL_BG)
        panel.pack(fill="both", expand=True)

        section_label(panel, "BATCH SETTINGS").pack(anchor="w", padx=14, pady=(14, 4))

        self._batch_folder = tk.StringVar()
        self._batch_out_folder = tk.StringVar()

        for lbl, var, cmd in [
            ("Source folder",  self._batch_folder,     self._browse_batch_folder),
            ("Output folder",  self._batch_out_folder, self._browse_batch_out_folder),
        ]:
            entry_row(panel, lbl, var, cmd).pack(fill="x", padx=14, pady=3)

        self._batch_custom_lbl = tk.Label(panel, text="", bg=PANEL_BG, fg=TEXT_DIM,
                                           font=("Segoe UI", 8), wraplength=300, justify="left")
        self._batch_custom_lbl.pack(anchor="w", padx=14, pady=(2, 0))

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)
        section_label(panel, "DETECTION SETTINGS").pack(anchor="w", padx=14, pady=(0, 6))

        conf_row = tk.Frame(panel, bg=PANEL_BG)
        conf_row.pack(fill="x", padx=14, pady=(0, 4))
        tk.Label(conf_row, text="Min. confidence", bg=PANEL_BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9), width=18, anchor="w").pack(side="left")
        self._batch_confidence = tk.DoubleVar(value=DEFAULT_CONFIDENCE)
        tk.Scale(conf_row, from_=0.3, to=1.0, resolution=0.05, orient="horizontal",
                 variable=self._batch_confidence, bg=PANEL_BG, fg=TEXT_MAIN,
                 troughcolor=ENTRY_BG, highlightthickness=0,
                 font=("Segoe UI", 8), activebackground=ACCENT, length=120
                 ).pack(side="left")
        self._batch_conf_lbl = tk.Label(conf_row, text=f"{DEFAULT_CONFIDENCE:.2f}",
                                         bg=PANEL_BG, fg=ACCENT, font=("Segoe UI", 9, "bold"), width=4)
        self._batch_conf_lbl.pack(side="left")
        self._batch_confidence.trace_add("write",
            lambda *_: self._batch_conf_lbl.config(text=f"{self._batch_confidence.get():.2f}"))

        tk.Frame(panel, bg=BORDER, height=1).pack(fill="x", padx=14, pady=10)

        info_box = tk.Frame(panel, bg=ENTRY_BG)
        info_box.pack(fill="x", padx=14, pady=(0, 10))
        tk.Label(info_box,
                 text="Batch mode processes every .txt, .docx, and .pdf\n"
                      "file in the source folder.\n\n"
                      "Each file gets its own anonymized output and\n"
                      "mapping file in the output folder.\n\n"
                      "If a _custom_pii.json file exists in the source\n"
                      "folder, it is automatically applied to all files.",
                 bg=ENTRY_BG, fg=TEXT_DIM, font=("Segoe UI", 9),
                 justify="left", padx=12, pady=10).pack()

        styled_button(panel, "\U0001f4c2  Run Batch Anonymize", self._run_batch,
                      width=28).pack(padx=14, pady=(0, 14))

        # Right: log output
        section_label(right, "BATCH LOG").pack(anchor="w", padx=4, pady=(4, 6))
        self._batch_log = scrolledtext.ScrolledText(
            right, bg=ENTRY_BG, fg=TEXT_MAIN, insertbackground=TEXT_MAIN,
            font=("Consolas", 9), relief="flat", wrap="word",
            selectbackground=ACCENT, selectforeground="white", state="disabled")
        self._batch_log.pack(fill="both", expand=True)
        self._batch_log.tag_configure("ok",    foreground=SUCCESS)
        self._batch_log.tag_configure("err",   foreground=DANGER)
        self._batch_log.tag_configure("info",  foreground=ACCENT)
        self._batch_log.tag_configure("warn",  foreground=WARNING)

    def _browse_batch_folder(self):
        folder = filedialog.askdirectory(title="Select source folder")
        if folder:
            self._batch_folder.set(folder)
            # Auto-set output folder to a subfolder
            self._batch_out_folder.set(os.path.join(folder, "anonymized"))
            # Check for project custom PII
            pii_path = get_project_pii_path(folder)
            if os.path.exists(pii_path):
                entries = load_project_pii(folder)
                self._batch_custom_lbl.config(
                    text=f"\u2714 Custom PII loaded: {len(entries)} entries from {CUSTOM_PII_FILENAME}",
                    fg=SUCCESS)
                # Also sync to the custom PII editor
                self._load_project_pii_into_editor(folder)
            else:
                self._batch_custom_lbl.config(
                    text=f"No {CUSTOM_PII_FILENAME} found — using NLP only",
                    fg=TEXT_DIM)

    def _browse_batch_out_folder(self):
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self._batch_out_folder.set(folder)

    def _batch_log_write(self, msg: str, tag: str = ""):
        self._batch_log.config(state="normal")
        self._batch_log.insert("end", msg + "\n", tag)
        self._batch_log.see("end")
        self._batch_log.config(state="disabled")

    def _run_batch(self):
        if not self._engine_ready:
            messagebox.showwarning("Engine Loading", "Please wait for the NLP engine to finish loading.")
            return
        src = self._batch_folder.get().strip()
        out = self._batch_out_folder.get().strip()
        if not src or not os.path.isdir(src):
            messagebox.showwarning("Missing Folder", "Please select a valid source folder.")
            return
        if not out:
            messagebox.showwarning("Missing Output", "Please specify an output folder.")
            return

        # Collect files
        exts = (".txt", ".docx", ".pdf")
        files = [f for f in os.listdir(src)
                 if os.path.splitext(f)[1].lower() in exts
                 and not f.startswith("_")]
        if not files:
            messagebox.showinfo("No Files", f"No .txt, .docx, or .pdf files found in:\n{src}")
            return

        confidence = self._batch_confidence.get()
        custom_entries = self._get_custom_pii_entries()
        # Also check for project-level custom PII in the source folder
        folder_entries = load_project_pii(src)
        all_custom = {e["text"]: e for e in (folder_entries + custom_entries)}
        merged_custom = list(all_custom.values())

        os.makedirs(out, exist_ok=True)

        # Clear log
        self._batch_log.config(state="normal")
        self._batch_log.delete("1.0", "end")
        self._batch_log.config(state="disabled")

        self._batch_log_write(f"Batch anonymize: {len(files)} file(s) in {src}", "info")
        if merged_custom:
            self._batch_log_write(f"Custom PII entries: {len(merged_custom)}", "info")
        self._batch_log_write("", "")

        self._start_spinner(f"Batch processing {len(files)} files...")
        self._set_status(f"Batch processing {len(files)} files...", WARNING)

        logger.info("Batch anonymize started: %d files in %s", len(files), src)
        if merged_custom:
            logger.info("Batch: %d custom PII entries loaded", len(merged_custom))

        def task():
            ok = 0
            fail = 0
            for fname in files:
                fpath = os.path.join(src, fname)
                base  = os.path.splitext(fname)[0]
                out_path = os.path.join(out, base + "_anonymized.txt")
                map_path = os.path.join(out, base + "_mapping.json")
                try:
                    logger.info("Batch: processing %s", fname)
                    text = read_file(fpath)
                    engine = PIIEngine.get()
                    anon_text, mapping, detections = engine.anonymize(
                        text, confidence=confidence)
                    if merged_custom:
                        entity_counts: Dict[str, int] = {}
                        value_to_ph: Dict[str, str] = {ph: orig for ph, orig in mapping.items()}
                        for ph in mapping:
                            m2 = re.match(r'\{\{([A-Z_]+)_(\d+)\}\}', ph)
                            if m2:
                                lbl2, num2 = m2.group(1), int(m2.group(2))
                                entity_counts[lbl2] = max(entity_counts.get(lbl2, 0), num2)
                        anon_text = apply_custom_pii(
                            anon_text, merged_custom,
                            mapping, entity_counts, value_to_ph, detections)
                    write_file(out_path, anon_text)
                    with open(map_path, "w", encoding="utf-8") as fh:
                        json.dump(mapping, fh, indent=4, ensure_ascii=False)
                    # Also write translation map for batch files
                    try:
                        trans_map_path = _get_translation_map_path(map_path)
                        trans_map = generate_translation_map(mapping)
                        with open(trans_map_path, "w", encoding="utf-8") as fh:
                            json.dump(trans_map, fh, indent=4, ensure_ascii=False)
                    except Exception as exc:
                        logger.warning("Batch: could not write translation map for %s: %s", fname, exc)
                    n = len(mapping)
                    logger.info("Batch: OK  %s  -> %d PII items", fname, n)
                    self.after(0, lambda fn=fname, n=n:
                        self._batch_log_write(f"  \u2714  {fn}  \u2192  {n} PII items replaced", "ok"))
                    ok += 1
                except Exception as exc:
                    logger.error("Batch: FAIL  %s  -> %s", fname, exc, exc_info=True)
                    self.after(0, lambda fn=fname, e=str(exc):
                        self._batch_log_write(f"  \u2718  {fn}  \u2014  ERROR: {e}", "err"))
                    fail += 1

            def done():
                logger.info("Batch complete: %d succeeded, %d failed", ok, fail)
                self._stop_spinner()
                self._batch_log_write("", "")
                self._batch_log_write(
                    f"Batch complete: {ok} succeeded, {fail} failed.", "info")
                self._batch_log_write(f"Output folder: {out}", "info")
                self._set_status(
                    f"Batch done: {ok} succeeded, {fail} failed.",
                    SUCCESS if fail == 0 else WARNING)
                messagebox.showinfo("Batch Complete",
                    f"Batch anonymization complete!\n\n"
                    f"  Succeeded : {ok}\n"
                    f"  Failed    : {fail}\n\n"
                    f"Output folder:\n  {out}")
            self.after(0, done)

        threading.Thread(target=task, daemon=True).start()


# ---------------------------------------------------------------------------
#  Entry point
# ---------------------------------------------------------------------------

def main():
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()
